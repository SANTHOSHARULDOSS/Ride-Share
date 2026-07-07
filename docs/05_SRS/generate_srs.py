import os
import sys
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

# Define styles and colors for Word document (Minimalist/Monochrome Corporate Style)
HEADINGS_COLOR = RGBColor(34, 34, 34)    # Dark Charcoal
BODY_COLOR = RGBColor(51, 51, 51)        # Charcoal
MUTED_COLOR = RGBColor(102, 102, 102)    # Medium Slate Gray
LIGHT_BG_HEX = "F2F2F2"                  # Minimalist Light Gray for table headers
BORDER_HEX = "CCCCCC"                    # Soft Gray for borders

def set_cell_background(cell, fill_hex):
    """Sets background color of a table cell."""
    tc_pr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tc_pr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell padding (in dxas)."""
    tc_pr = cell._element.get_or_add_tcPr()
    tc_mar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tc_mar.append(node)
    tc_pr.append(tc_mar)

def format_row(row, bg_hex=None, bold=False, font_size=10):
    """Formats all cells in a row using Calibri font."""
    for cell in row.cells:
        if bg_hex:
            set_cell_background(cell, bg_hex)
        set_cell_margins(cell)
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(font_size)
                run.font.bold = bold
                run.font.color.rgb = BODY_COLOR

def add_heading_styled(doc, text, level):
    """Adds a heading with Cambria styling."""
    p = doc.add_heading('', level=level)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    run.font.name = 'Cambria'
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = HEADINGS_COLOR
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = HEADINGS_COLOR
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = HEADINGS_COLOR
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = HEADINGS_COLOR
    return p

def add_paragraph_styled(doc, text="", bold_prefix=None, bullet=False):
    """Adds a paragraph with Calibri styling."""
    style_name = 'List Bullet' if bullet else 'Normal'
    p = doc.add_paragraph(style=style_name)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    
    if bold_prefix:
        r_prefix = p.add_run(bold_prefix)
        r_prefix.font.name = 'Calibri'
        r_prefix.font.bold = True
        r_prefix.font.size = Pt(11)
        r_prefix.font.color.rgb = BODY_COLOR
        
    if text:
        r_text = p.add_run(text)
        r_text.font.name = 'Calibri'
        r_text.font.size = Pt(11)
        r_text.font.color.rgb = BODY_COLOR
    return p

def add_code_styled(doc, code_text):
    """Adds a block of preformatted monospace code."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.25)
    
    pPr = p._element.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
    pPr.append(shd)
    
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9.5)
    run.font.color.rgb = BODY_COLOR
    return p

def add_image_styled(doc, image_name, caption_text):
    """Adds a centered image with a styled caption below it."""
    image_path = os.path.join(r"d:\Ride-Share\screenshots", image_name)
    if not os.path.exists(image_path):
        print(f"Warning: Image not found at {image_path}. Skipping.")
        return None
        
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(10)
    p_img.paragraph_format.space_after = Pt(4)
    
    run_img = p_img.add_run()
    try:
        run_img.add_picture(image_path, width=Inches(5.5))
    except Exception as e:
        print(f"Error adding picture {image_name}: {e}")
        return None
        
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_before = Pt(2)
    p_cap.paragraph_format.space_after = Pt(12)
    p_cap.paragraph_format.keep_with_next = True
    
    run_cap = p_cap.add_run(caption_text)
    run_cap.font.name = 'Calibri'
    run_cap.font.size = Pt(9.5)
    run_cap.font.italic = True
    run_cap.font.color.rgb = MUTED_COLOR
    return p_img

def build_docx(md_sections, filepath):
    """Generates a professional Word Document based on structured sections."""
    doc = Document()
    
    # Page setup - Standard 1 inch margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Set default styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = BODY_COLOR
    
    # ------------------ COVER PAGE (Cambria/Calibri Minimal Style) ------------------
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(120)
    p_title.paragraph_format.space_after = Pt(12)
    r_title = p_title.add_run("AI-Based Smart Route Ride-Sharing System\n")
    r_title.font.name = 'Cambria'
    r_title.font.size = Pt(28)
    r_title.font.bold = True
    r_title.font.color.rgb = HEADINGS_COLOR
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(80)
    r_sub = p_sub.add_run("Software Requirements Specification (SRS) Document\n")
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(16)
    r_sub.font.color.rgb = MUTED_COLOR
    
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_after = Pt(120)
    
    metadata = [
        ("Version", "1.0.0 (Production-Inspired MVP)"),
        ("Document Status", "Approved for Implementation"),
        ("Prepared By", "=== FILLED BY USER ==="),
        ("Organization", "=== FILLED BY USER ==="),
        ("Client", "=== FILLED BY USER ==="),
        ("Date", "July 7, 2026"),
        ("Confidentiality", "=== FILLED BY USER ==="),
        ("Classification", "=== FILLED BY USER ===")
    ]
    for key, val in metadata:
        r_lbl = p_meta.add_run(f"{key}: ")
        r_lbl.font.name = 'Calibri'
        r_lbl.font.bold = True
        r_lbl.font.color.rgb = BODY_COLOR
        
        r_val = p_meta.add_run(f"{val}\n")
        r_val.font.name = 'Calibri'
        r_val.font.color.rgb = BODY_COLOR
    
    # Page Break after Cover Page
    doc.add_page_break()
    
    # ------------------ DOCUMENT CONTROL ------------------
    add_heading_styled(doc, "Document Control", 1)
    
    add_paragraph_styled(doc, "Document Revision History", bold_prefix="Revision History:")
    t_rev = doc.add_table(rows=3, cols=5)
    t_rev.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_rev.autofit = False
    
    # Column widths
    col_widths = [Inches(0.8), Inches(1.2), Inches(1.0), Inches(1.5), Inches(2.0)]
    for row in t_rev.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width
            
    # Headers
    t_rev.rows[0].cells[0].text = "Ver."
    t_rev.rows[0].cells[1].text = "Date"
    t_rev.rows[0].cells[2].text = "Author"
    t_rev.rows[0].cells[3].text = "Reviewer(s)"
    t_rev.rows[0].cells[4].text = "Description of Change"
    format_row(t_rev.rows[0], LIGHT_BG_HEX, bold=True)
    
    # Row 1
    t_rev.rows[1].cells[0].text = "0.1.0"
    t_rev.rows[1].cells[1].text = "2026-07-01"
    t_rev.rows[1].cells[2].text = "=== FILLED BY USER ==="
    t_rev.rows[1].cells[3].text = "=== FILLED BY USER ==="
    t_rev.rows[1].cells[4].text = "Initial draft outline & requirements gathering."
    format_row(t_rev.rows[1])
    
    # Row 2
    t_rev.rows[2].cells[0].text = "1.0.0"
    t_rev.rows[2].cells[1].text = "2026-07-07"
    t_rev.rows[2].cells[2].text = "=== FILLED BY USER ==="
    t_rev.rows[2].cells[3].text = "=== FILLED BY USER ==="
    t_rev.rows[2].cells[4].text = "Baseline release mapping fully to source models and features."
    format_row(t_rev.rows[2])
    
    add_paragraph_styled(doc, "")
    add_paragraph_styled(doc, "Document Approvers", bold_prefix="Approvals:")
    t_app = doc.add_table(rows=2, cols=4)
    t_app.alignment = WD_TABLE_ALIGNMENT.CENTER
    app_widths = [Inches(1.8), Inches(1.8), Inches(1.4), Inches(1.5)]
    for row in t_app.rows:
        for idx, w in enumerate(app_widths):
            row.cells[idx].width = w
            
    t_app.rows[0].cells[0].text = "Name"
    t_app.rows[0].cells[1].text = "Role"
    t_app.rows[0].cells[2].text = "Signature"
    t_app.rows[0].cells[3].text = "Date"
    format_row(t_app.rows[0], LIGHT_BG_HEX, bold=True)
    
    t_app.rows[1].cells[0].text = "=== FILLED BY USER ==="
    t_app.rows[1].cells[1].text = "=== FILLED BY USER ==="
    t_app.rows[1].cells[2].text = "=== FILLED BY USER ==="
    t_app.rows[1].cells[3].text = "=== FILLED BY USER ==="
    format_row(t_app.rows[1])
    
    add_paragraph_styled(doc, "")
    add_paragraph_styled(doc, "Distribution List", bold_prefix="Distribution:")
    t_dist = doc.add_table(rows=2, cols=3)
    t_dist.alignment = WD_TABLE_ALIGNMENT.CENTER
    dist_widths = [Inches(2.2), Inches(2.2), Inches(2.1)]
    for row in t_dist.rows:
        for idx, w in enumerate(dist_widths):
            row.cells[idx].width = w
            
    t_dist.rows[0].cells[0].text = "Name / Group"
    t_dist.rows[0].cells[1].text = "Organization / Role"
    t_dist.rows[0].cells[2].text = "Date Shared"
    format_row(t_dist.rows[0], LIGHT_BG_HEX, bold=True)
    
    t_dist.rows[1].cells[0].text = "=== FILLED BY USER ==="
    t_dist.rows[1].cells[1].text = "=== FILLED BY USER ==="
    t_dist.rows[1].cells[2].text = "=== FILLED BY USER ==="
    format_row(t_dist.rows[1])
    
    doc.add_page_break()
    
    # ------------------ TABLE OF CONTENTS ------------------
    add_heading_styled(doc, "Table of Contents", 1)
    
    toc_items = [
        ("1. Introduction", 1),
        ("1.1 Purpose", 2), ("1.2 Scope", 2), ("1.3 Intended Audience", 2), ("1.4 Definitions", 2),
        ("1.5 Acronyms", 2), ("1.6 Abbreviations", 2), ("1.7 References", 2), ("1.8 Business Context", 2),
        ("1.9 Background", 2), ("1.10 Project Vision", 2), ("1.11 Objectives", 2), ("1.12 Success Criteria", 2),
        ("2. Overall Description", 1),
        ("2.1 Product Perspective", 2), ("2.2 Product Positioning", 2), ("2.3 Business Goals", 2),
        ("2.4 Product Features", 2), ("2.5 User Classes", 2), ("2.6 User Characteristics", 2),
        ("2.7 Operating Environment", 2), ("2.8 Development Environment", 2), ("2.9 Technology Stack", 2),
        ("2.10 Design Constraints", 2), ("2.11 Assumptions", 2), ("2.12 Dependencies", 2),
        ("3. Stakeholders", 1),
        ("4. Functional Requirements", 1),
        ("5. Non-Functional Requirements", 1),
        ("6. User Roles & Permissions Matrix", 1),
        ("7. User Journey Maps", 1),
        ("8. Complete System Modules", 1),
        ("9. Use Cases", 1),
        ("10. User Stories", 1),
        ("11. Workflows (ASCII Diagrams)", 1),
        ("12. Database Design", 1),
        ("12.1 Entity Relationship Diagram", 2), ("12.2 Database Schema Tables", 2), ("12.3 Normalization", 2),
        ("13. API Documentation", 1),
        ("14. AI Module Specifications", 1),
        ("15. Email System", 1),
        ("16. Security Infrastructure", 1),
        ("17. UI/UX Design System", 1),
        ("18. External Integrations", 1),
        ("19. Deployment Guide", 1),
        ("20. Testing & Verification", 1),
        ("21. Risks Assessment", 1),
        ("22. Future Enhancements", 1),
        ("23. Maintenance Strategy", 1),
        ("24. Project Timeline", 1),
        ("25. Budget & Cost Breakdown", 1),
        ("26. Legal & Compliance", 1),
        ("27. Appendices", 1)
    ]
    
    for label, level in toc_items:
        p_toc = doc.add_paragraph()
        p_toc.paragraph_format.space_before = Pt(1)
        p_toc.paragraph_format.space_after = Pt(1)
        indent = Inches(0.25 * (level - 1))
        p_toc.paragraph_format.left_indent = indent
        
        run = p_toc.add_run(label)
        run.font.name = 'Calibri'
        run.font.size = Pt(11 if level == 1 else 9.5)
        run.font.bold = (level == 1)
        run.font.color.rgb = BODY_COLOR
        
    doc.add_page_break()
    
    # ------------------ WRITING MAIN SECTIONS ------------------
    for key, sec_content in md_sections.items():
        # Title of section
        title = sec_content.get('title', key)
        level = sec_content.get('level', 1)
        add_heading_styled(doc, title, level)
        
        # Content blocks
        blocks = sec_content.get('blocks', [])
        for block in blocks:
            b_type = block.get('type', 'p')
            text = block.get('text', '')
            bold_prefix = block.get('bold_prefix', None)
            
            if b_type == 'p':
                add_paragraph_styled(doc, text, bold_prefix=bold_prefix)
            elif b_type == 'bullet':
                add_paragraph_styled(doc, text, bold_prefix=bold_prefix, bullet=True)
            elif b_type == 'code':
                add_code_styled(doc, text)
            elif b_type == 'image':
                add_image_styled(doc, block.get('name'), text)
            elif b_type == 'table':
                # Table rendering
                headers = block.get('headers', [])
                rows_data = block.get('rows', [])
                col_widths_inches = block.get('widths', [])
                
                table = doc.add_table(rows=len(rows_data) + 1, cols=len(headers))
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.autofit = False
                
                # Apply widths
                for row in table.rows:
                    for c_idx, w in enumerate(col_widths_inches):
                        row.cells[c_idx].width = Inches(w)
                        
                # Headers text
                for c_idx, h in enumerate(headers):
                    table.rows[0].cells[c_idx].text = str(h)
                format_row(table.rows[0], LIGHT_BG_HEX, bold=True, font_size=9)
                
                # Rows text
                for r_idx, row_text in enumerate(rows_data):
                    for c_idx, val in enumerate(row_text):
                        table.rows[r_idx + 1].cells[c_idx].text = str(val)
                    format_row(table.rows[r_idx + 1], font_size=8.5)
                
                # Spacer
                add_paragraph_styled(doc, "")
                
    doc.save(filepath)
    print(f"Word document saved to {filepath}")

# ----------------- MAIN DATA STRUCTURE DEFINITION -----------------
srs_sections = {}

srs_sections["1.0"] = {
    "title": "1. Introduction",
    "level": 1,
    "blocks": [
        {"type": "p", "text": "This document defines the Software Requirements Specification (SRS) for the AI-Based Smart Route Ride-Sharing System (herein referred to as 'Ride-Share'). It lays out the comprehensive functional, non-functional, security, and architectural specifications that govern the implementation and verification of the platform's production MVP."},
        {"type": "p", "bold_prefix": "1.1 Purpose: ", "text": "The primary objective of this document is to establish a rigorous, industry-grade baseline requirements specification. This serves as the ground truth for developers, quality assurance engineers, administrators, and university/organizational compliance evaluators. A software development team can directly initiate database schema setups, API routing, frontend interface coding, and system testing from this document."},
        {"type": "p", "bold_prefix": "1.2 Scope: ", "text": "The Ride-Share system is an offline-capable Progressive Web Application (PWA) facilitating path-sequence carpooling matches. Rather than basic start-to-end taxi matches, Ride-Share incorporates geodetic route waypoint sequences (via Leaflet.js maps) to identify drivers and passengers traveling on overlapping routes. It handles secure role-based registration, vehicle verification, automated match rankings based on mathematical proximity detours and trust relationships (shared corporate/college organizations, friendships, event attendance), real-time instant messaging with attachment handling, simulated checkout and payments, a SOS panic telemetry system, and an AI-driven automated support framework."},
        {"type": "p", "bold_prefix": "1.3 Intended Audience: ", "text": "This document is targeted at backend/frontend engineers, QA test practitioners, DevOps deployment teams, evaluators, and system managers. Reading this document requires intermediate knowledge of Django MVC architectures, REST APIs, databases, geodetic route sequencing, and Web App service caching systems."},
        {"type": "p", "bold_prefix": "1.4 Definitions: ", "text": "Terms used in this document are defined as follows:"},
        {"type": "bullet", "bold_prefix": "Driver: ", "text": "A registered traveler who owns a vehicle, publishes carpooling offers with sequential geographic stop waypoints, and charges passenger fares to offset transit expenses."},
        {"type": "bullet", "bold_prefix": "Passenger: ", "text": "A traveler who searches for ride listings matching their sequential pickup/drop points, requests bookings, and performs simulated payments."},
        {"type": "bullet", "bold_prefix": "Waypoint: ", "text": "An ordered set of GPS coordinates (latitude, longitude) along a driver's route representing stops or pick-up/drop-off intervals."},
        {"type": "bullet", "bold_prefix": "Detour Proximity: ", "text": "The total deviation distance in kilometers that a driver must execute to accommodate a passenger's pickup and drop-off points."},
        {"type": "bullet", "bold_prefix": "Social Affiliation: ", "text": "Shared context metrics (common employer, college, friends) used to calculate safety and trust ratings."},
        {"type": "p", "bold_prefix": "1.5 Acronyms: ", "text": "The following acronyms are used throughout this document:"},
        {"type": "bullet", "text": "SRS: Software Requirements Specification"},
        {"type": "bullet", "text": "PWA: Progressive Web Application"},
        {"type": "bullet", "text": "API: Application Programming Interface"},
        {"type": "bullet", "text": "OSM: OpenStreetMap"},
        {"type": "bullet", "text": "OSRM: Open Source Routing Machine"},
        {"type": "bullet", "text": "JWT: JSON Web Token"},
        {"type": "bullet", "text": "CSRF: Cross-Site Request Forgery"},
        {"type": "bullet", "text": "XSS: Cross-Site Scripting"},
        {"type": "bullet", "text": "CORS: Cross-Origin Resource Sharing"},
        {"type": "bullet", "text": "GDPR: General Data Protection Regulation"},
        {"type": "bullet", "text": "SMTP: Simple Mail Transfer Protocol"},
        {"type": "p", "bold_prefix": "1.6 Abbreviations: ", "text": "The following abbreviations are utilized: 'Ver.' (Version), 'Temp.' (Template), 'Notif.' (Notification), 'Org' (Organization), 'DB' (Database), 'Sec.' (Section), 'Req.' (Requirement)."},
        {"type": "p", "bold_prefix": "1.7 References: ", "text": "Key reference documents and standards include:"},
        {"type": "bullet", "text": "IEEE Std 830-1998: Recommended Practice for Software Requirements Specifications."},
        {"type": "bullet", "text": "ISO/IEC/IEEE 29148:2018: Systems and software engineering - Lifecycle processes - Requirements engineering."},
        {"type": "bullet", "text": "Django 6.0 Web Framework Documentation (https://docs.djangoproject.com/)"},
        {"type": "bullet", "text": "Leaflet.js Mapping Library API (https://leafletjs.com/)"},
        {"type": "bullet", "text": "Google Gemini API Documentation (https://ai.google.dev/)"},
        {"type": "bullet", "text": "Ride-Share Active Database models (core/models.py) and view logic (core/views.py, core/views_extended.py)."},
        {"type": "p", "bold_prefix": "1.8 Business Context: ", "text": "Urban centers face massive traffic congestions, escalating vehicle ownership costs, and carbon footprints. Conventional commercial taxi aggregators charge steep surge fares while taking large commissions. A community-based peer-to-peer carpooling platform enables commuters to offset fuel costs while helping passengers commute affordably in a secure, trust-validated circle."},
        {"type": "p", "bold_prefix": "1.9 Background: ", "text": "The project was designed as a high-fidelity academic prototype showing clean software engineering principles. It integrates geodetic calculations, web caching, maps, and AI modules into a single lightweight Django package to prove technical capability without high infrastructure billing overheads."},
        {"type": "p", "bold_prefix": "1.10 Project Vision: ", "text": "To deliver a seamless, offline-first, geodetically intelligent carpooling platform that eliminates search friction, minimizes detours, maximizes personal safety, and handles inquiries through advanced AI integrations."},
        {"type": "p", "bold_prefix": "1.11 Objectives: ", "text": "1. Provide sequential route matching calculations with low latencies. 2. Enable installable, offline-first dashboards to handle connectivity interruptions. 3. Establish a trust network combining identity document vetting and social circle metrics. 4. Offer an AI helper chatbot and travel itinerary builder."},
        {"type": "p", "bold_prefix": "1.12 Success Criteria: ", "text": "1. All unit tests execute successfully. 2. 100% database seeding via custom commands. 3. Offline dashboard renders when network is lost. 4. Geodetic matching yields correct sequence outputs."}
    ]
}

srs_sections["2.0"] = {
    "title": "2. Overall Description",
    "level": 1,
    "blocks": [
        {"type": "p", "bold_prefix": "2.1 Product Perspective: ", "text": "Ride-Share is a standalone Progressive Web Application running a Django 6 backend and HTML5/Bootstrap client interface. It integrates with OpenStreetMap API via Leaflet.js to sequence route coordinates, utilizes the local database (SQLite/PostgreSQL) for persistence, and interacts with Google's Gemini-1.5-flash LLM model to offer conversational customer helpdesk support and route itinerary suggestions."},
        {"type": "p", "bold_prefix": "2.2 Product Positioning: ", "text": "Positions itself as a private, high-trust, cost-sharing carpooling solution for universities, corporations, and structured communities, bridging the gap between informal messaging group arrangements and commercial taxi services."},
        {"type": "p", "bold_prefix": "2.3 Business Goals: ", "text": "Reduce daily commuting expenses by dividing vehicle fuel costs equally among passengers, and decrease the number of single-occupancy vehicles on key arterial corridors."},
        {"type": "p", "bold_prefix": "2.4 Product Features: ", "text": "The primary capabilities include: 1. Sequential Waypoint Matching Engine. 2. Leaflet Map Editor. 3. Mock Real-time Vehicle Tracking. 4. SOS Emergency Alarms. 5. Social Communities (College, Org, Friends). 6. Event RSVPs. 7. Real-Time Chat with File Upload. 8. Verification & Reputation. 9. AI Assistants (Chatbot, Trip Planner, support reply auto-classifier). 10. Service Worker PWA Offline mode."},
        {"type": "p", "bold_prefix": "2.5 User Classes: ", "text": "The system identifies three primary user categories: Passengers, Drivers, and Administrators."},
        {"type": "p", "bold_prefix": "2.6 User Characteristics: ", "text": "Users are assumed to be daily commuters (students or office workers) possessing a mobile device or laptop with an internet browser. Users require zero training to navigate the interface due to intuitive, standard web UI structures."},
        {"type": "p", "bold_prefix": "2.7 Operating Environment: ", "text": "Client-side: Any modern web browser (Chrome, Safari, Firefox, Edge). Can be installed as a PWA locally. Server-side: Hosted on standard Python 3.10+ runtimes on Windows or Linux, using SQLite for development/testing and PostgreSQL for production environments."},
        {"type": "p", "bold_prefix": "2.8 Development Environment: ", "text": "Local Windows/OSX/Linux machine, Python virtual environment, Visual Studio Code / Cursor IDE, git version control, and django CLI tools."},
        {"type": "p", "bold_prefix": "2.9 Technology Stack: ", "text": "The platform is constructed on: Backend: Django 6.0, Django Channels (WebSockets). Frontend: HTML5, CSS3 Custom Properties (Light/Dark themes), Bootstrap 5, JavaScript (ES6). Maps: Leaflet.js with OpenStreetMap tiles, Nominatim API, OSRM. Database: SQLite (local file) / PostgreSQL. AI: Google Gemini-1.5-flash API via Google Generative AI SDK."},
        {"type": "p", "bold_prefix": "2.10 Design Constraints: ", "text": "Must run on low-bandwidth connections. The route-matching algorithm must not use external paid APIs (e.g. Google Maps API) and must perform calculations locally. Service worker caching must adhere to PWA standards. Database writes must occur safely in SQLite without concurrency deadlocks."},
        {"type": "p", "bold_prefix": "2.11 Assumptions: ", "text": "1. **Assumption:** The host browser supports ES6 and Service Workers. 2. **Assumption:** The network allows HTTPS connections to external OpenStreetMap servers. 3. **Assumption:** Seeding coordinates resemble genuine roadways to ensure math logic succeeds."},
        {"type": "p", "bold_prefix": "2.12 Dependencies: ", "text": "Availability of OSM mapping server layers, OSRM router server for path calculation, and internet connectivity to communicate with Google's Gemini LLM API."}
    ]
}

srs_sections["3.0"] = {
    "title": "3. Stakeholders",
    "level": 1,
    "blocks": [
        {"type": "p", "text": "The stakeholders involved in the life-cycle and operation of the Ride-Share platform are classified below:"},
        {"type": "table", "headers": ["Stakeholder", "Description / Role", "Primary Concerns / Requirements"],
         "widths": [1.5, 2.5, 2.5],
         "rows": [
             ["Project Owner", "=== FILLED BY USER ===", "Ensures project timelines are met, budget compliance, and overall business value."],
             ["Client", "=== FILLED BY USER ===", "Acceptance testing, deployment validation, and core feature alignment."],
             ["End Users - Passenger", "Commuters seeking affordable transit.", "Accurate search matches, low fare costs, verification trust, and emergency SOS alerts."],
             ["End Users - Driver", "Vehicle owners seeking cost offset.", "Simple route publishing, seating capacity control, passenger verification, and safe navigation."],
             ["Administrators", "Operational managers and monitors.", "User verification approval, system metrics monitoring, and handling safety reports/support tickets."],
             ["Developers", "Software engineering team.", "Clean codebase, fast local seeding, standard API paths, and robust testing frameworks."],
             ["Testers", "Quality Assurance team.", "Comprehensive unit/integration testing, security vulnerabilities checks, and cross-browser testing."],
             ["Support Team", "Customer helpdesk staff.", "Access to support tickets list, AI-assisted auto replies, and dashboard controls."],
             ["Management", "=== FILLED BY USER ===", "Strategic decisions, marketing, and scaling plans."],
             ["External Services", "Leaflet, OSM, Gemini APIs.", "API stability, rate-limiting quotas, response times, and connection security."],
             ["Third-party Providers", "=== FILLED BY USER ===", "Hosting servers (Render, Vercel), email SMTP relays, and domain name registrars."]
         ]}
    ]
}

srs_sections["4.0"] = {
    "title": "4. Functional Requirements",
    "level": 1,
    "blocks": [
        {"type": "p", "text": "This section details the primary functional requirements of the Ride-Share platform. All features are categorized with unique tracking IDs, input/output paths, priorities, and acceptance criteria."},
        
        {"type": "p", "bold_prefix": "FR-001: Account Registration", "text": ""},
        {"type": "bullet", "bold_prefix": "Description: ", "text": "Allows a user to create a new profile as a Passenger or Driver."},
        {"type": "bullet", "bold_prefix": "Priority: ", "text": "High"},
        {"type": "bullet", "bold_prefix": "Inputs: ", "text": "Username, email, password, role choice (Driver/Passenger), phone number."},
        {"type": "bullet", "bold_prefix": "Outputs: ", "text": "Database user record, email verification token, redirection to login page."},
        {"type": "bullet", "bold_prefix": "Preconditions: ", "text": "The username must be unique; email must not already exist in DB."},
        {"type": "bullet", "bold_prefix": "Postconditions: ", "text": "A User record is saved with verified flags set to False. An activation email is logged/simulated."},
        {"type": "bullet", "bold_prefix": "Acceptance Criteria: ", "text": "The system validates password strength, triggers checking of username uniqueness via AJAX helper, and stores the user role choice correctly."},
        {"type": "bullet", "bold_prefix": "Dependencies: ", "text": "Database engine, Email verification system."},
        
        {"type": "p", "bold_prefix": "FR-002: User Authentication (Login / Logout)", "text": ""},
        {"type": "bullet", "bold_prefix": "Description: ", "text": "Authenticates returning users, issues secure session cookies, and records login history (UserSession model)."},
        {"type": "bullet", "bold_prefix": "Priority: ", "text": "High"},
        {"type": "bullet", "bold_prefix": "Inputs: ", "text": "Username or Email, Password."},
        {"type": "bullet", "bold_prefix": "Outputs: ", "text": "Session ID, UserSession record in DB, redirection to role-specific dashboard."},
        {"type": "bullet", "bold_prefix": "Preconditions: ", "text": "User must possess an active, registered account. Verification of email is recommended but not blocking for MVP login."},
        {"type": "bullet", "bold_prefix": "Postconditions: ", "text": "User is authenticated; a UserSession model registers IP address, user agent, and device type."},
        {"type": "bullet", "bold_prefix": "Acceptance Criteria: ", "text": "Invalid credentials prompt error messages. Successful login creates a session cookie. Redirection lands on Passenger, Driver, or Admin panel depending on User.role."},
        {"type": "bullet", "bold_prefix": "Dependencies: ", "text": "Django Auth framework, Database engine."},
        
        {"type": "p", "bold_prefix": "FR-003: Vehicle Details Management", "text": ""},
        {"type": "bullet", "bold_prefix": "Description: ", "text": "Allows a Driver to register and update their vehicle properties."},
        {"type": "bullet", "bold_prefix": "Priority: ", "text": "High"},
        {"type": "bullet", "bold_prefix": "Inputs: ", "text": "Make, model, license plate, seating capacity, color."},
        {"type": "bullet", "bold_prefix": "Outputs: ", "text": "Vehicle record associated with the Driver's profile in DB."},
        {"type": "bullet", "bold_prefix": "Preconditions: ", "text": "Authenticated user must have the role of 'DRIVER'."},
        {"type": "bullet", "bold_prefix": "Postconditions: ", "text": "Vehicle record is created; verified status set to False until admin review."},
        {"type": "bullet", "bold_prefix": "Acceptance Criteria: ", "text": "The license plate must be unique. Seating capacity must be a positive integer. Vehicle details are accessible on profile fields."},
        {"type": "bullet", "bold_prefix": "Dependencies: ", "text": "FR-002, DB Engine."},
        
        {"type": "p", "bold_prefix": "FR-004: Publish Ride Offer", "text": ""},
        {"type": "bullet", "bold_prefix": "Description: ", "text": "Enables a Driver to schedule a carpooling offer, selecting route waypoints on an interactive map editor."},
        {"type": "bullet", "bold_prefix": "Priority: ", "text": "High"},
        {"type": "bullet", "bold_prefix": "Inputs: ", "text": "Vehicle selection, start location address, end location address, departure time, seat price, available seats, cost sharing terms, luggage details, emergency contact, list of ordered map waypoint coordinates."},
        {"type": "bullet", "bold_prefix": "Outputs: ", "text": "Ride record, sequential RouteWaypoint records linked via ForeignKey, redirection to active list."},
        {"type": "bullet", "bold_prefix": "Preconditions: ", "text": "Driver is logged in and has at least one registered vehicle in the system."},
        {"type": "bullet", "bold_prefix": "Postconditions: ", "text": "Ride record is saved with status 'ACTIVE' (or 'DRAFT'). RouteWaypoint coordinates sequence is indexed into DB."},
        {"type": "bullet", "bold_prefix": "Acceptance Criteria: ", "text": "Waypoints are sequenced on the Leaflet map and stored in sequence_order. Available seats must be less than or equal to the selected vehicle's capacity. Price must be positive."},
        {"type": "bullet", "bold_prefix": "Dependencies: ", "text": "FR-003, Leaflet/OSM map scripts."},
        
        {"type": "p", "bold_prefix": "FR-005: Intelligent Geodetic Route Waypoint Matching Search", "text": ""},
        {"type": "bullet", "bold_prefix": "Description: ", "text": "Allows passengers to query routes. The backend computes matches by scanning sequential waypoints and calculating detour distances using the Haversine formula, returning ranked recommendations with social trust boosts."},
        {"type": "bullet", "bold_prefix": "Priority: ", "text": "High"},
        {"type": "bullet", "bold_prefix": "Inputs: ", "text": "Pickup location coordinates, dropoff location coordinates, seats needed, maximum walk distance (km)."},
        {"type": "bullet", "bold_prefix": "Outputs: ", "text": "Ordered JSON list of matching Ride models, pickup/dropoff walk distances, base scores, trust boosts, match reasons, and final recommendation scores."},
        {"type": "bullet", "bold_prefix": "Preconditions: ", "text": "Passenger is logged in. Active rides must exist in the database with departure time >= current time."},
        {"type": "bullet", "bold_prefix": "Postconditions: ", "text": "No data is modified. Read-only recommendation scoring list is returned to user interface."},
        {"type": "bullet", "bold_prefix": "Acceptance Criteria: ", "text": "Matches are filtered to ensure pickup waypoint sequence index is less than dropoff waypoint sequence index (guarantees heading in the same direction). Matches are sorted with recommendation score descending (Base Proximity + Trust boosts). Walk distance must be <= maximum walk detour constraint."},
        {"type": "bullet", "bold_prefix": "Dependencies: ", "text": "route_matching.py logic, SQLite spatial decimals."},
        
        {"type": "p", "bold_prefix": "FR-006: Create Ride Booking", "text": ""},
        {"type": "bullet", "bold_prefix": "Description: ", "text": "Enables a Passenger to book seats on a matched ride route."},
        {"type": "bullet", "bold_prefix": "Priority: ", "text": "High"},
        {"type": "bullet", "bold_prefix": "Inputs: ", "text": "Ride ID, pickup location, dropoff location, pickup lat/lng, dropoff lat/lng, seats requested."},
        {"type": "bullet", "bold_prefix": "Outputs: ", "text": "Booking record in DB with status 'PENDING', fare estimation total, notification to driver."},
        {"type": "bullet", "bold_prefix": "Preconditions: ", "text": "Passenger is logged in. Selected ride has available seats >= requested seats."},
        {"type": "bullet", "bold_prefix": "Postconditions: ", "text": "Booking is registered in DB; driver is sent a Notification model alert (type 'TRIP_REQUEST')."},
        {"type": "bullet", "bold_prefix": "Acceptance Criteria: ", "text": "The total price is calculated as (seats_requested * ride.price_per_seat). System validates that driver is not booking their own ride."},
        {"type": "bullet", "bold_prefix": "Dependencies: ", "text": "FR-004, Notification model."},
        
        {"type": "p", "bold_prefix": "FR-007: Booking Approval Lifecycle", "text": ""},
        {"type": "bullet", "bold_prefix": "Description: ", "text": "Allows a Driver to accept or reject pending bookings, modifying seats availability."},
        {"type": "bullet", "bold_prefix": "Priority: ", "text": "High"},
        {"type": "bullet", "bold_prefix": "Inputs: ", "text": "Booking ID, Action choice (Accept / Reject)."},
        {"type": "bullet", "bold_prefix": "Outputs: ", "text": "Updated Booking record status, Notification model alert to passenger, modified Ride available_seats count."},
        {"type": "bullet", "bold_prefix": "Preconditions: ", "text": "The logged-in user must be the Driver who published the corresponding Ride."},
        {"type": "bullet", "bold_prefix": "Postconditions: ", "text": "Booking status transitions to 'ACCEPTED' or 'REJECTED'. If accepted, Ride.available_seats decrements by Booking.seats_requested."},
        {"type": "bullet", "bold_prefix": "Acceptance Criteria: ", "text": "Declining a booking releases seats. Accepting a booking automatically checks if seats capacity is exceeded, updating DB transactionally. Redirection logs confirmation."},
        {"type": "bullet", "bold_prefix": "Dependencies: ", "text": "FR-006, Notification module."},
        
        {"type": "p", "bold_prefix": "FR-008: Real-Time Chat & File Attachment", "text": ""},
        {"type": "bullet", "bold_prefix": "Description: ", "text": "Provides chat rooms for rides, communities, or direct users, allowing file attachment uploads."},
        {"type": "bullet", "bold_prefix": "Priority: ", "text": "Medium"},
        {"type": "bullet", "bold_prefix": "Inputs: ", "text": "Room type (RIDE/COMMUNITY/DIRECT), Room ID, text content, optional file payload."},
        {"type": "bullet", "bold_prefix": "Outputs: ", "text": "Message record in DB, live message broadcast (Django Channels Websocket or polling fallback)."},
        {"type": "bullet", "bold_prefix": "Preconditions: ", "text": "User is a member of the community, active passenger/driver in the ride, or verified friend of direct contact."},
        {"type": "bullet", "bold_prefix": "Postconditions: ", "text": "Message model is stored; notification of unread message is flagged for recipient."},
        {"type": "bullet", "bold_prefix": "Acceptance Criteria: ", "text": "Attachments are checked for safety. File size limit is 5MB. Messages are rendered in chronological order. HTML tag sanitization is applied."},
        {"type": "bullet", "bold_prefix": "Dependencies: ", "text": "Message model, media upload directory."},
        
        {"type": "p", "bold_prefix": "FR-009: Social Communities Creation & Member Moderation", "text": ""},
        {"type": "bullet", "bold_prefix": "Description: ", "text": "Allows users to create groups based on domains (college or employer) to filter ride circles."},
        {"type": "bullet", "bold_prefix": "Priority: ", "text": "Medium"},
        {"type": "bullet", "bold_prefix": "Inputs: ", "text": "Community name, category, description, type (Public/Private), avatar/banner graphics, joining rules."},
        {"type": "bullet", "bold_prefix": "Outputs: ", "text": "Community record, CommunityMember record with role 'ADMIN' for the creator."},
        {"type": "bullet", "bold_prefix": "Preconditions: ", "text": "Creator is an authenticated user with verified credentials (if creating private/domain groups)."},
        {"type": "bullet", "bold_prefix": "Postconditions: ", "text": "Community model added. Members can query and join."},
        {"type": "bullet", "bold_prefix": "Acceptance Criteria: ", "text": "Domain emails match domain rules. Joining request status defaults to 'PENDING' for private groups, requiring Community Admin approval action."},
        {"type": "bullet", "bold_prefix": "Dependencies: ", "text": "Community, CommunityMember models."},
        
        {"type": "p", "bold_prefix": "FR-010: Event Scheduling & RSVPs", "text": ""},
        {"type": "bullet", "bold_prefix": "Description: ", "text": "Provides events creation inside communities, enabling carpooling planning to a common venue (e.g. hackathons, trek meetups)."},
        {"type": "bullet", "bold_prefix": "Priority: ", "text": "Medium"},
        {"type": "bullet", "bold_prefix": "Inputs: ", "text": "Title, description, location, date, event type, community selection."},
        {"type": "bullet", "bold_prefix": "Outputs: ", "text": "Event record, EventAttendee record with status 'GOING' for creator."},
        {"type": "bullet", "bold_prefix": "Preconditions: ", "text": "User is a member of the chosen community."},
        {"type": "bullet", "bold_prefix": "Postconditions: ", "text": "Event details listed; community members can post RSVP status choice (GOING/INTERESTED)."},
        {"type": "bullet", "bold_prefix": "Acceptance Criteria: ", "text": "RSVPs update attendee records instantly. Users attending the same event receive a +10 trust boost in route search."},
        {"type": "bullet", "bold_prefix": "Dependencies: ", "text": "Event, EventAttendee models."},
        
        {"type": "p", "bold_prefix": "FR-011: Simulated SOS Emergency Alerts", "text": ""},
        {"type": "bullet", "bold_prefix": "Description: ", "text": "Enables a passenger or driver in danger to trigger an alarm, generating live telemetry and sounding a local warning siren."},
        {"type": "bullet", "bold_prefix": "Priority: ", "text": "High"},
        {"type": "bullet", "bold_prefix": "Inputs: ", "text": "Ride ID, current coordinates (latitude, longitude)."},
        {"type": "bullet", "bold_prefix": "Outputs: ", "text": "SOSAlert record in DB, active status notification, audible sound activation on frontend client, admin logs alert."},
        {"type": "bullet", "bold_prefix": "Preconditions: ", "text": "The user is an active passenger or driver on a trip with status 'ACTIVE'."},
        {"type": "bullet", "bold_prefix": "Postconditions: ", "text": "An active SOSAlert record is stored. Emergency contact fields in Ride are logged for escalation."},
        {"type": "bullet", "bold_prefix": "Acceptance Criteria: ", "text": "Audible alert triggers immediately. Latency to save the SOSAlert record must be under 50ms. Admin panel displays active alarms in flashing red."},
        {"type": "bullet", "bold_prefix": "Dependencies: ", "text": "SOSAlert model, Frontend JS audio element."},
        
        {"type": "p", "bold_prefix": "FR-012: ID and Domain Email Verification", "text": ""},
        {"type": "bullet", "bold_prefix": "Description: ", "text": "Validates user trust levels using ID image uploads and corporate/college email verification tokens."},
        {"type": "bullet", "bold_prefix": "Priority: ", "text": "Medium"},
        {"type": "bullet", "bold_prefix": "Inputs: ", "text": "File upload (Gov ID card) or domain email input (e.g. student@college.edu)."},
        {"type": "bullet", "bold_prefix": "Outputs: ", "text": "Database verification token, profile status changes, email verification log."},
        {"type": "bullet", "bold_prefix": "Preconditions: ", "text": "Authenticated user is logged in."},
        {"type": "bullet", "bold_prefix": "Postconditions: ", "text": "is_email_verified or is_gov_id_verified set to True upon token entry or admin review."},
        {"type": "bullet", "bold_prefix": "Acceptance Criteria: ", "text": "Vetting token expires in 24 hours. Verification email is simulated with a custom link containing the token parameter. Correct token entry updates status."},
        {"type": "bullet", "bold_prefix": "Dependencies: ", "text": "EmailVerificationToken model, Email dispatch framework."},
        
        {"type": "p", "bold_prefix": "FR-013: AI Support Chatbot Help Desk", "text": ""},
        {"type": "bullet", "bold_prefix": "Description: ", "text": "Integrates a 24/7 support assistant powered by Gemini-1.5-flash to reply to platform FAQs, safety procedures, and community questions."},
        {"type": "bullet", "bold_prefix": "Priority: ", "text": "Medium"},
        {"type": "bullet", "bold_prefix": "Inputs: ", "text": "User question, chat history text."},
        {"type": "bullet", "bold_prefix": "Outputs: ", "text": "AI response string, AILog record in database."},
        {"type": "bullet", "bold_prefix": "Preconditions: ", "text": "User is authenticated and logged in."},
        {"type": "bullet", "bold_prefix": "Postconditions: ", "text": "Question and reply are written to the AILog table with log_type set to 'CHATBOT'."},
        {"type": "bullet", "bold_prefix": "Acceptance Criteria: ", "text": "If Gemini API key is missing, system falls back to predefined string matching (e.g. if 'sos' is queried, returns standard emergency instructions). Responses are formatted in markdown."},
        {"type": "bullet", "bold_prefix": "Dependencies: ", "text": "Gemini API, AILog model, ai_services.py."},
        
        {"type": "p", "bold_prefix": "FR-014: AI Travel Itinerary Planner", "text": ""},
        {"type": "bullet", "bold_prefix": "Description: ", "text": "Uses Gemini to construct travel itineraries between source and destination endpoints, detailing route advice, local stops, weather forecasts, and packing lists."},
        {"type": "bullet", "bold_prefix": "Priority: ", "text": "Medium"},
        {"type": "bullet", "bold_prefix": "Inputs: ", "text": "Source address, destination address, travel preferences description."},
        {"type": "bullet", "bold_prefix": "Outputs: ", "text": "Markdown-formatted itinerary page, AILog entry."},
        {"type": "bullet", "bold_prefix": "Preconditions: ", "text": "User is authenticated."},
        {"type": "bullet", "bold_prefix": "Postconditions: ", "text": "Generates AILog entry with type 'TRAVEL_ASSISTANT'."},
        {"type": "bullet", "bold_prefix": "Acceptance Criteria: ", "text": "If API fails, a formatted high-quality mock response containing NH route details is rendered. Formatted chapters match the standard output headings (Estimated Route, Hotels, Weather, packing checklists)."},
        {"type": "bullet", "bold_prefix": "Dependencies: ", "text": "Gemini API, AILog model."},
        
        {"type": "p", "bold_prefix": "FR-015: Support Ticket Creation & AI Auto-Reply Classifier", "text": ""},
        {"type": "bullet", "bold_prefix": "Description: ", "text": "Allows visitors to submit support tickets. The backend processes the email contents through Gemini to detect the category, draft an auto-response, and determine if admin escalation is required."},
        {"type": "bullet", "bold_prefix": "Priority: ", "text": "Medium"},
        {"type": "bullet", "bold_prefix": "Inputs: ", "text": "Name, email, subject, message body."},
        {"type": "bullet", "bold_prefix": "Outputs: ", "text": "SupportTicket record, AI classification parameters, AILog, notification email details."},
        {"type": "bullet", "bold_prefix": "Preconditions: ", "text": "None (Accessible on public contact view)."},
        {"type": "bullet", "bold_prefix": "Postconditions: ", "text": "A SupportTicket record is stored. Ticket status set to 'OPEN'. Needs_escalation boolean is saved. AI draft auto-reply is saved in ticket fields."},
        {"type": "bullet", "bold_prefix": "Acceptance Criteria: ", "text": "The ticket is assigned a unique alphanumeric ticket number (e.g. RS-XXXXX). Emails containing safety threats or abuse trigger needs_escalation=True automatically. Admin dashboard shows classified tickets sorted by status and urgency."},
        {"type": "bullet", "bold_prefix": "Dependencies: ", "text": "SupportTicket model, Gemini API classifier, ai_services.py."}
    ]
}

srs_sections["5.0"] = {
    "title": "5. Non-Functional Requirements",
    "level": 1,
    "blocks": [
        {"type": "p", "text": "This section outlines the non-functional requirements (NFRs) that specify the system's operational standards, security baselines, and quality attributes."},
        {"type": "bullet", "bold_prefix": "Performance: ", "text": "The route-matching algorithm (Haversine sequential checks) must compute matches in less than 100ms for database sizes up to 10,000 active rides. Main dashboard pages must render in under 1.5 seconds under ordinary network latencies."},
        {"type": "bullet", "bold_prefix": "Availability: ", "text": "Target system availability is 99.9% uptime (excluding scheduled maintenance windows, which must occur during off-peak hours 2:00 AM - 4:00 AM UTC)."},
        {"type": "bullet", "bold_prefix": "Reliability: ", "text": "The Mean Time Between Failures (MTBF) must be greater than 500 operating hours. Data corruption rate must remain at 0% by executing atomic transactions on database bookings."},
        {"type": "bullet", "bold_prefix": "Scalability: ", "text": "The application architecture separates static assets (served via Whitenoise or CDN) from computational routes. For production, the database engine must transition from SQLite to PostgreSQL to support connection pooling and up to 5,000 concurrent user sessions."},
        {"type": "bullet", "bold_prefix": "Maintainability: ", "text": "The backend must strictly adhere to Django MVC architecture and PEP-8 coding style guidelines. Code changes must achieve a minimum of 80% unit test coverage before deployment approval."},
        {"type": "bullet", "bold_prefix": "Security: ", "text": "All user passwords must be hashed using PBKDF2 with SHA-256. Cross-Site Request Forgery (CSRF) tokens must be validated on all POST/PUT requests. Rate limiting (django-ratelimit) must restrict brute-force login attempts to a maximum of 5 tries per IP within 5 minutes (django-axes)."},
        {"type": "bullet", "bold_prefix": "Privacy: ", "text": "Personally Identifiable Information (PII) including phone numbers, government ID documents, and geolocations must be encrypted at rest. Users must be able to delete their account, executing cascading deletion of all personal profiles in compliance with GDPR guidelines."},
        {"type": "bullet", "bold_prefix": "Accessibility: ", "text": "All HTML user interface elements must conform to Web Content Accessibility Guidelines (WCAG) 2.1 Level AA, including contrast ratios of 4.5:1, semantic page structures, and keyboard navigation capability."},
        {"type": "bullet", "bold_prefix": "Compatibility: ", "text": "Must run seamlessly across different devices and operating systems. Supported browsers: Google Chrome (version 90+), Apple Safari (version 14+), Mozilla Firefox (version 88+), and Microsoft Edge (version 90+)."},
        {"type": "bullet", "bold_prefix": "Localization: ", "text": "The system's primary interface language is English. Code structure must wrap all user-facing strings in translation hooks (Django ugettext) to support seamless future localization. UTC is used as the database timezone, converting to client local timezone in front-end views."},
        {"type": "bullet", "bold_prefix": "Portability: ", "text": "The system must run inside standard Docker containers to facilitate quick deployment to cloud service providers (Render, AWS, GCP, Heroku)."},
        {"type": "bullet", "bold_prefix": "Logging: ", "text": "Every API transaction, authentication lifecycle event, error trace, email dispatch, and AI prompt request must be logged. AI calls are persisted in the AILog table; emails are logged in the EmailLog table. General application errors are written to rotating files on the server (max 10MB per file)."},
        {"type": "bullet", "bold_prefix": "Monitoring: ", "text": "A public health check endpoint (/offline/ or /admin/status/) must return system status metrics. System monitoring tools (e.g. Sentry or Prometheus) must alert admins of elevated 5xx error responses within 60 seconds."},
        {"type": "bullet", "bold_prefix": "Backup: ", "text": "Daily automated backups of the SQLite file / PostgreSQL DB must be archived to secure cloud storage (e.g. AWS S3) with a retention policy of 30 days."},
        {"type": "bullet", "bold_prefix": "Disaster Recovery: ", "text": "In the event of database failure or server corruption, the Recovery Time Objective (RTO) must be less than 4 hours, and the Recovery Point Objective (RPO) must be less than 24 hours."},
        {"type": "bullet", "bold_prefix": "Fault Tolerance: ", "text": "If external services (Leaflet tiles, Gemini API, or SMTP server) fail, the core system must continue operating. Map pages must show warning banners, Gemini requests must fallback to local mock logic, and failed emails must queue for retry instead of throwing 500 errors to the client."},
        {"type": "bullet", "bold_prefix": "Compliance: ", "text": "The platform must adhere to local transport regulations regarding non-commercial carpooling (fares must only offset vehicle operating costs and not generate commercial profits)."},
        {"type": "bullet", "bold_prefix": "Browser Support: ", "text": "Must test and verify layout alignment and Javascript responsiveness on Chrome, Safari, Firefox, and Edge desktop versions."},
        {"type": "bullet", "bold_prefix": "Mobile Support: ", "text": "Must support installable PWA manifest structures on Android (Chrome PWA launcher) and iOS (Add to Home Screen Safari shortcut) with responsive layouts adapting down to 320px viewport widths."}
    ]
}

srs_sections["6.0"] = {
    "title": "6. User Roles & Permissions Matrix",
    "level": 1,
    "blocks": [
        {"type": "p", "text": "The system enforces rigid role-based access control (RBAC). The following matrix defines the exact actions permitted for each user category:"},
        {"type": "table", "headers": ["Feature Area", "Passenger", "Driver", "Administrator"],
         "widths": [2.2, 1.4, 1.4, 1.5],
         "rows": [
             ["Register & Profile Edit", "Allowed", "Allowed", "Allowed"],
             ["Register Vehicle", "Denied", "Allowed", "Allowed"],
             ["Publish Ride Offer", "Denied", "Allowed", "Allowed"],
             ["Search Ride Matches", "Allowed", "Allowed (as passenger)", "Allowed"],
             ["Create Booking Request", "Allowed", "Allowed (as passenger)", "Allowed"],
             ["Approve/Reject Booking", "Denied", "Allowed (own ride)", "Allowed (all)"],
             ["Trigger SOS Alarm", "Allowed (during trip)", "Allowed (during trip)", "Allowed"],
             ["Create Community", "Allowed", "Allowed", "Allowed"],
             ["Approve Community Member", "Denied", "Denied (unless admin)", "Allowed (all)"],
             ["Post in Community Feed", "Allowed", "Allowed", "Allowed"],
             ["Real-time Chat Rooms", "Allowed (own rides/comms)", "Allowed (own rides/comms)", "Allowed (all)"],
             ["Send Friend Request", "Allowed", "Allowed", "Allowed"],
             ["Submit Safety Report", "Allowed", "Allowed", "Allowed"],
             ["Use AI Chatbot / Planner", "Allowed", "Allowed", "Allowed"],
             ["Access System Logs", "Denied", "Denied", "Allowed"],
             ["Reset System Data", "Denied", "Denied", "Allowed"],
             ["Vetting ID Uploads", "Denied", "Denied", "Allowed"]
         ]}
    ]
}

srs_sections["7.0"] = {
    "title": "7. User Journey Maps",
    "level": 1,
    "blocks": [
        {"type": "p", "text": "This section traces the chronological user journey maps across the lifecycle of the application interfaces:"},
        {"type": "bullet", "bold_prefix": "Registration: ", "text": "1. User lands on landing page. 2. Clicks 'Register'. 3. Fills registration fields. 4. Submits form. 5. System checks username uniqueness via AJAX. 6. DB stores inactive record. 7. Activation token generated and simulated. 8. Redirected to email verification page."},
        {"type": "image", "name": "registration_page.png", "text": "Figure 7.1: User Registration Form View"},
        {"type": "bullet", "bold_prefix": "Login: ", "text": "1. User visits /login/. 2. Enters credentials or clicks single-click quick-access demo profile buttons (Admin, Passenger, Driver). 3. System checks PBKDF2 hash. 4. Records UserSession metadata. 5. Redirects to dashboard."},
        {"type": "image", "name": "login_page.png", "text": "Figure 7.2: User Authentication Form View"},
        {"type": "bullet", "bold_prefix": "Forgot Password: ", "text": "1. User clicks 'Forgot Password' on login page. 2. Enters registered email. 3. System saves PasswordResetToken (expires in 1 hour). 4. Simulates email dispatch. 5. User clicks link. 6. Enters new password on reset form. 7. Password hashes and updates in DB. 8. Session tokens expire."},
        {"type": "bullet", "bold_prefix": "Dashboard Access: ", "text": "1. Logs in. 2. Role check executes. 3. Passenger dashboard displays active bookings, search bars, and community feeds. 4. Driver dashboard displays vehicle status, active ride offers, pending booking requests list, and trip stats. 5. Admin dashboard lists audit logs, safety reports, and AI usage summaries."},
        {"type": "bullet", "bold_prefix": "Feature Access (Ride Booking): ", "text": "1. Passenger inputs source and destination. 2. Filters matches by distance detour. 3. Views recommendation scores. 4. Clicks 'Book Ride'. 5. Selects seat quantity. 6. Confirm Booking. 7. System forwards trip request to Driver. 8. Driver clicks 'Approve'. 9. Passenger status changes to accepted, prompting payment checkout page."},
        {"type": "bullet", "bold_prefix": "Notifications: ", "text": "1. User receives instant popup or dashboard badge. 2. Navigates to notifications view. 3. Reviews list of request events. 4. Clicks 'Mark Read' or 'Mark All Read'. 5. DB updates is_read flag. 6. UI badge count resets."},
        {"type": "bullet", "bold_prefix": "Profile Management: ", "text": "1. User navigates to /profile/. 2. Updates bio, skills, interests, and profile pictures. 3. If Driver, registers car model, capacity, and plate. 4. Submits organizational email. 5. User gets college/corporate verification token to verify domain affiliation. 6. Submits ID card scan for admin review."},
        {"type": "bullet", "bold_prefix": "Logout: ", "text": "1. User clicks 'Logout' in header. 2. Django backend destroys session. 3. UserSession.is_active set to False. 4. Client clear session caches. 5. Redirected to landing page with connection status shown."},
        {"type": "bullet", "bold_prefix": "Account Deletion: ", "text": "1. User selects 'Delete Account' in profile settings. 2. Confirms password verification challenge. 3. System performs cascade deletes of User, vehicles, bookings, friendships, and messages. 4. Logs out user. 5. Audits account removal."}
    ]
}

srs_sections["8.0"] = {
    "title": "8. Complete System Modules",
    "level": 1,
    "blocks": [
        {"type": "p", "text": "The architecture is split into separate modules that govern specific operations within the Django MVC environment:"},
        {"type": "bullet", "bold_prefix": "Authentication & Authorization Modules: ", "text": "Extends AbstractUser to custom User model. Handles registration form validation, email token checks, session token tracking, quick-login seeding, and Google OAuth bindings. Provides route protection decorators (@login_required, custom role verification middleware) to block unauthorized operations."},
        {"type": "bullet", "bold_prefix": "Dashboard Module: ", "text": "Serves three specialized HTML views: Passenger UI, Driver UI, and Admin UI. Coordinates metrics, listings, and workflows dynamically."},
        {"type": "image", "name": "passenger_dashboard.png", "text": "Figure 8.1: Passenger Portal Dashboard Interface"},
        {"type": "image", "name": "driver_dashboard.png", "text": "Figure 8.2: Driver Commute Portal Dashboard View"},
        {"type": "image", "name": "admin_dashboard.png", "text": "Figure 8.3: Administrator Analytics Dashboard View"},
        {"type": "bullet", "bold_prefix": "Profile & Settings Module: ", "text": "Manages bio descriptions, skill tags, languages, travel rules (smoking, pets), and vehicle models using custom forms validated in backend core/forms.py."},
        {"type": "image", "name": "profile_page.png", "text": "Figure 8.4: User Profile Customization and Validation Page"},
        {"type": "bullet", "bold_prefix": "Search Module: ", "text": "Funnels pickup and dropoff coordinates to the geodetic route sequencer, returning compatible rides ranked by distance detour scores."},
        {"type": "image", "name": "ride_search.png", "text": "Figure 8.5: Geodetic Waypoint Routing Matching Search Results"},
        {"type": "bullet", "bold_prefix": "Community Module: ", "text": "Manages user community circles, private group invites, approval loops, posts, comments, announcements, and feed likes."},
        {"type": "image", "name": "communities.png", "text": "Figure 8.6: Community Circles Directory View"},
        {"type": "bullet", "bold_prefix": "Notifications Module: ", "text": "Implements the Notification model. Dispatches real-time alerts upon booking approvals, friend request events, and community invitations."},
        {"type": "image", "name": "notifications.png", "text": "Figure 8.7: Notifications Alerts Center Console"},
        {"type": "bullet", "bold_prefix": "Logging & Audit Module: ", "text": "Captures system actions. Persists all generative AI prompts/responses in AILog, and outgoing emails in EmailLog tables."},
        {"type": "image", "name": "admin_system_logs.png", "text": "Figure 8.8: Admin System Security & Session Log Registry"}
    ]
}

srs_sections["9.0"] = {
    "title": "9. Use Cases",
    "level": 1,
    "blocks": [
        {"type": "p", "text": "The core use cases detailing system workflows are defined below:"},
        
        {"type": "p", "bold_prefix": "Use Case ID: UC-01: Publish Route-Based Ride Offer", "text": ""},
        {"type": "bullet", "bold_prefix": "Actor: ", "text": "Driver"},
        {"type": "bullet", "bold_prefix": "Description: ", "text": "Driver outlines their commute path on Leaflet, schedules the ride details, and publishes the offer to the matching pool."},
        {"type": "bullet", "bold_prefix": "Preconditions: ", "text": "Driver is authenticated, holds verified vehicle details in DB, and has active status."},
        {"type": "bullet", "bold_prefix": "Main Flow: ", "text": "1. Driver navigates to 'Publish Ride'. 2. Driver selects their vehicle. 3. Map interface loads. 4. Driver enters Start and End locations. 5. Driver clicks on intermediate streets to pin waypoint coordinates. 6. Enters date, departure time, seats capacity, cost terms, luggage rules. 7. Submits form. 8. System validates inputs. 9. Saves Ride and RouteWaypoint models. 10. Redirects to dashboard with success message."},
        {"type": "bullet", "bold_prefix": "Alternative Flow (Empty Waypoints): ", "text": "At Step 5, if no intermediate waypoints are pinned, system calculates standard straight path sequence containing start and end coordinates as waypoints, then proceeds."},
        {"type": "bullet", "bold_prefix": "Exception Flow (Outdated Departure): ", "text": "At Step 6, if departure time is in the past, system throws validation error on form, highlighting time field, and prompts corrected entry."},
        {"type": "bullet", "bold_prefix": "Postconditions: ", "text": "Ride record created with status set to ACTIVE. Waypoint coordinates are sequenced in DB."},
        {"type": "bullet", "bold_prefix": "Business Rules: ", "text": "Fare charges cannot exceed maximum local cost-sharing guidelines. Vehicle capacity constraint must match registration."},
        {"type": "bullet", "bold_prefix": "Acceptance Criteria: ", "text": "Published ride is searchable by pickup/drop locations. Map displays route waypoints correctly in order."},
        
        {"type": "p", "bold_prefix": "Use Case ID: UC-02: Search & Book Overlapping Route", "text": ""},
        {"type": "bullet", "bold_prefix": "Actor: ", "text": "Passenger"},
        {"type": "bullet", "bold_prefix": "Description: ", "text": "Passenger enters pickup and drop points, runs geodetic proximity scan, views recommendation scoring, and requests seats booking."},
        {"type": "bullet", "bold_prefix": "Preconditions: ", "text": "Passenger is authenticated, has active status, and is not the driver of matching rides."},
        {"type": "bullet", "bold_prefix": "Main Flow: ", "text": "1. Passenger navigates to 'Search Rides'. 2. Enters pickup address and dropoff address. 3. System geocodes addresses to coordinates. 4. Runs haversine distance matching query on active routes waypoints list. 5. Backend ranks matching rides based on detour walk distance and social trust affiliations. 6. Passenger reviews recommended ride card. 7. Clicks 'Book Ride'. 8. Enters seats count. 9. Confirms. 10. Booking created with status 'PENDING'. 11. Redirection logs confirmation."},
        {"type": "bullet", "bold_prefix": "Alternative Flow (No Direct Matches): ", "text": "At Step 5, if detour walk distance exceeds max limit (e.g. 5km), system reports 'No rides found', suggests increasing walk threshold or modifying search locations."},
        {"type": "bullet", "bold_prefix": "Exception Flow (Seats Exceeded): ", "text": "At Step 8, if seats requested exceeds available_seats on Ride, system raises error, block submission, and prompts user to reduce quantity."},
        {"type": "bullet", "bold_prefix": "Postconditions: ", "text": "Booking record is added. Notification model registers alert for Driver."},
        {"type": "bullet", "bold_prefix": "Business Rules: ", "text": "Passengers cannot book seats on expired rides. Maximum detour walking distance defaults to 5.0 km."},
        {"type": "bullet", "bold_prefix": "Acceptance Criteria: ", "text": "Matching calculations correctly identify route sequences (pickup waypoint index must be chronologically prior to dropoff waypoint index along driver route)."},
        
        {"type": "p", "bold_prefix": "Use Case ID: UC-03: Trigger SOS Emergency Alarm", "text": ""},
        {"type": "bullet", "bold_prefix": "Actor: ", "text": "Passenger or Driver"},
        {"type": "bullet", "bold_prefix": "Description: ", "text": "Triggers local warning siren and updates database telemetry alerts for administrator escalation."},
        {"type": "bullet", "bold_prefix": "Preconditions: ", "text": "User is active driver or passenger on a ride with status 'ACTIVE'."},
        {"type": "bullet", "bold_prefix": "Main Flow: ", "text": "1. User navigates to tracking page. 2. Clicks red 'SOS' button. 3. System prompts 'Confirm Emergency?'. 4. User clicks confirm. 5. Frontend JS triggers audible warning sound loop on client browser. 6. Transmits current GPS coordinate parameters. 7. Backend saves SOSAlert record. 8. Admin log shows flash red alarm. 9. Simulated emergency contact alert triggers."},
        {"type": "bullet", "bold_prefix": "Alternative Flow (No GPS Coordinates): ", "text": "At Step 6, if browser GPS permission is denied, system reads last known waypoint coordinates of ride record, saves alert, and proceeds."},
        {"type": "bullet", "bold_prefix": "Exception Flow (Expired Ride): ", "text": "At Step 1, if trip is already status 'COMPLETED', SOS button is disabled on UI, blocking trigger."},
        {"type": "bullet", "bold_prefix": "Postconditions: ", "text": "SOSAlert record is stored as is_active=True. Administrative alarms active."},
        {"type": "bullet", "bold_prefix": "Business Rules: ", "text": "SOS triggers must bypass ordinary API rate limiting constraints."},
        {"type": "bullet", "bold_prefix": "Acceptance Criteria: ", "text": "Siren audio plays locally. SOS logs are saved successfully in database. Admin dashboard displays emergency state immediately."},
        
        {"type": "p", "bold_prefix": "Use Case ID: UC-04: Auto-Reply Support Ticket", "text": ""},
        {"type": "bullet", "bold_prefix": "Actor: ", "text": "Visitor / Customer"},
        {"type": "bullet", "bold_prefix": "Description: ", "text": "Processes incoming customer emails, classifies issue types with Gemini AI, saves auto-reply draft, and flags escalations."},
        {"type": "bullet", "bold_prefix": "Preconditions: ", "text": "None (Accessible publicly on contact view)."},
        {"type": "bullet", "bold_prefix": "Main Flow: ", "text": "1. User visits contact page. 2. Fills Name, Email, Subject, and Message. 3. Clicks 'Submit Ticket'. 4. System creates SupportTicket record. 5. Backend forwards text parameters to reply_support_email function. 6. Gemini parses text, returning JSON classification structure. 7. System updates ticket category, records AI reply suggestions, sets needs_escalation flag. 8. Returns submission confirmation to user."},
        {"type": "bullet", "bold_prefix": "Alternative Flow (API Key Missing): ", "text": "At Step 5, if GEMINI_API_KEY is not defined, system triggers local regex filter, detects terms (like 'abuse', 'hack'), saves category, sets standard mock reply text, and completes."},
        {"type": "bullet", "bold_prefix": "Exception Flow (Corrupt JSON Payload): ", "text": "At Step 6, if Gemini returns invalid format, system logs JSON parse exception, sets category to 'General Inquiry', sets escalation to True (defaulting to safe review), and continues."},
        {"type": "bullet", "bold_prefix": "Postconditions: ", "text": "SupportTicket is added. EmailLog persists simulated auto-reply dispatch details."},
        {"type": "bullet", "bold_prefix": "Business Rules: ", "text": "Tickets marked with needs_escalation=True must appear at the top of the Admin support inbox queue."},
        {"type": "bullet", "bold_prefix": "Acceptance Criteria: ", "text": "Support tickets are generated with unique alphanumeric codes (RS-XXXXX). AI auto-reply drafts map to issue categories (e.g. Account Issue, abuse)."}
    ]
}

srs_sections["10.0"] = {
    "title": "10. User Stories",
    "level": 1,
    "blocks": [
        {"type": "p", "text": "The software development team must test and deliver against the following agile user stories:"},
        
        {"type": "p", "bold_prefix": "User Story US-001: Passenger Booking Search", "text": ""},
        {"type": "bullet", "bold_prefix": "Story: ", "text": "As a daily Passenger commuter, I want to search for rides by entering my pickup and drop addresses, so that I can find drivers who are already traveling on overlapping routes and avoid paying commercial taxi fares."},
        {"type": "bullet", "bold_prefix": "Acceptance Criteria 1: ", "text": "Search results must display matching rides where my pickup point is located before my dropoff point along the driver's route sequence."},
        {"type": "bullet", "bold_prefix": "Acceptance Criteria 2: ", "text": "Rides must display a detailed walkthrough of calculated walk distances and recommendation scores."},
        {"type": "bullet", "bold_prefix": "Acceptance Criteria 3: ", "text": "Inactive or full rides must be automatically filtered out of search listings."},
        
        {"type": "p", "bold_prefix": "User Story US-002: Driver Waypoints Map Editor", "text": ""},
        {"type": "bullet", "bold_prefix": "Story: ", "text": "As a commuting Driver, I want to pins intermediate waypoints on a map when publishing my ride, so that passengers at intermediate locations along my journey can match and join my carpool."},
        {"type": "bullet", "bold_prefix": "Acceptance Criteria 1: ", "text": "The UI must load an interactive Leaflet.js map with geocoding searches to easily set start and end points."},
        {"type": "bullet", "bold_prefix": "Acceptance Criteria 2: ", "text": "Driver clicks on map must append sequential waypoint coordinate values to a hidden array form field."},
        {"type": "bullet", "bold_prefix": "Acceptance Criteria 3: ", "text": "The route line must visually update to show the waypoint sequence connection on the map canvas."},
        
        {"type": "p", "bold_prefix": "User Story US-003: Safety Emergency SOS", "text": ""},
        {"type": "bullet", "bold_prefix": "Story: ", "text": "As a Passenger on a ride, I want to trigger a panic alarm with a single click in case of emergency, so that my driver is warned, local sirens alert bystanders, and safety admins log my exact coordinates for escalation."},
        {"type": "bullet", "bold_prefix": "Acceptance Criteria 1: ", "text": "The tracking dashboard must display a prominent, flashing Red 'SOS' button during active trips."},
        {"type": "bullet", "bold_prefix": "Acceptance Criteria 2: ", "text": "Clicking the button must play an audible warning siren loop on the device speaker immediately."},
        {"type": "bullet", "bold_prefix": "Acceptance Criteria 3: ", "text": "Backend must save the SOS alert with coordinates and set status to active within 50ms."},
        
        {"type": "p", "bold_prefix": "User Story US-004: Trust Badges Vetting", "text": ""},
        {"type": "bullet", "bold_prefix": "Story: ", "text": "As a User concerned with safety, I want to verify my corporate or university email address, so that I can unlock organization trust badges and find safe ride matches from my own institution."},
        {"type": "bullet", "bold_prefix": "Acceptance Criteria 1: ", "text": "Profile page must offer domain email submission fields."},
        {"type": "bullet", "bold_prefix": "Acceptance Criteria 2: ", "text": "Entering a domain email triggers a verification token code sent to that address."},
        {"type": "bullet", "bold_prefix": "Acceptance Criteria 3: ", "text": "Entering the correct token updates profile variables, displaying verified trust badges next to username."}
    ]
}

srs_sections["11.0"] = {
    "title": "11. Workflows (ASCII Diagrams)",
    "level": 1,
    "blocks": [
        {"type": "p", "text": "The operational sequence and activity workflows are diagrammed in ASCII formatting below:"},
        {"type": "p", "bold_prefix": "1. User Registration & Email Verification Lifecycle:", "text": ""},
        {"type": "code", "text": """
[User Client]                      [Django App Backend]                [Email Log/SMTP]
      |                                     |                                 |
      |--- 1. Submit Registration Form ---->|                                 |
      |    (User, Email, Password, Role)    |                                 |
      |                                     |--- 2. Hash Password (PBKDF2) -->|
      |                                     |--- 3. Create Inactive User ---->|
      |                                     |--- 4. Save Verification Token ->|
      |                                     |                                 |
      |<-- 5. Redirect to Verification -----|                                 |
      |                                     |--- 6. Dispatch simulated mail ->|
      |                                     |      (Contains Token link)      |
      |                                     |                                 |
      |--- 7. Click Token link (or input) ->|                                 |
      |                                     |--- 8. Validate Token expiry ----|
      |                                     |--- 9. Set is_email_verified=True|
      |<-- 10. Render Verified Profile -----|                                 |
        """},
        {"type": "p", "bold_prefix": "2. Ride Publishing, Waypoint Matching, and Booking Lifecycle:", "text": ""},
        {"type": "code", "text": """
[Driver Client]             [Maps API/OSRM]          [App Database]          [Passenger Client]
       |                           |                       |                          |
       |--- 1. Select Route ------>|                       |                          |
       |    (Pin Waypoints)        |                       |                          |
       |<-- 2. Draw Route Path ----|                       |                          |
       |                                                   |                          |
       |--- 3. Publish Ride (Time, Fare, Waypoints) ------>|                          |
       |                                                   |                          |
       |                                                   |--- 4. Search overlapping |
       |                                                   |      waypoints (pickup,  |
       |                                                   |      dropoff sequential) |
       |                                                   |                          |
       |                                                   |<-- 5. Match score list --|
       |                                                   |                          |
       |                                                   |<- 6. Submit Book request-|
       |<-- 7. Receive Pending Request Notification -------|                          |
       |                                                   |                          |
       |--- 8. Accept Booking Request -------------------->|                          |
       |                                                   |--- 9. Decrement seats ---|
       |                                                   |                          |
       |                                                   |<-- 10. Redirect checkout-|
       |                                                   |<-- 11. Simulate checkout-|
       |<-- 12. Trip Starts & Telemetry Tracks ------------|<-- 12. Trip Starts ------|
        """},
        {"type": "p", "bold_prefix": "3. SOS Panic Trigger Emergency Alert:", "text": ""},
        {"type": "code", "text": """
[User Tracking UI]               [Web API Client]               [Backend DB]           [Admin Panel]
        |                               |                            |                       |
        |--- 1. Clicks 'SOS Button' --->|                            |                       |
        |                               |--- 2. Get Geolocation ---->|                       |
        |--- 3. Siren sound loop (on) ->|                            |                       |
        |                               |--- 4. Dispatch Alert ----->|                       |
        |                               |      (Ride ID, Lat/Lng)    |                       |
        |                               |                            |--- 5. Save Alert ---->|
        |                               |                            |   (is_active=True)    |
        |                               |                            |                       |
        |                               |<-- 6. Trigger Admin alarm -|---------------------->|
        |                               |                            |                       | (Flash Red)
        |                               |<-- 7. Dispatch Contact list|                       |
        """},
        {"type": "p", "bold_prefix": "4. Automated Help Desk & Email Classifier (AI):", "text": ""},
        {"type": "code", "text": """
[Contact Page Visitor]           [Backend core]            [Gemini API Model]          [Admin Tickets]
         |                              |                           |                         |
         |--- 1. Submit Ticket ------->|                           |                         |
         |    (Name, Email, Message)    |                           |                         |
         |                              |--- 2. Forward contents -->|                         |
         |                              |      (Prompt Template)    |                         |
         |                              |                           |--- 3. Classify Category |
         |                              |                           |--- 4. Draft Auto-Reply  |
         |                              |                           |--- 5. Flag Escalation   |
         |                              |                           |                         |
         |                              |<-- 6. JSON payload -------|                         |
         |                              |                           |                         |
         |                              |--- 7. Save SupportTicket -|------------------------>|
         |<-- 8. Simulated Auto-Reply --|                           |                         | (Urgent Alert
                                                                                                 |  if escalated)
        """}
    ]
}

srs_sections["12.0"] = {
    "title": "12. Database Design",
    "level": 1,
    "blocks": [
        {"type": "p", "bold_prefix": "12.1 Entity Relationship Diagram (ASCII Layout):", "text": ""},
        {"type": "code", "text": """
  +------------------+                   +------------------+
  |       User       |1 ------------ N   |     Vehicle      |
  |------------------|                   |------------------|
  | PK: id           |                   | PK: id           |
  | username, role   |                   | FK: driver_id    |
  | reputation_points|                   | make, capacity   |
  +------------------+                   +------------------+
        |1        |1                              |1
        |         +---------------------+         |
        |                               v         v
        |                            +------------------+
        |                            |       Ride       |1 ----------+
        |                            |------------------|            |
        |                            | PK: id           |            |
        |                            | FK: driver_id    |            |
        |                            | FK: vehicle_id   |            |
        |                            | start_location   |            |
        |                            +------------------+            |
        |1                                 |1                        |N
        |                                  v                         v
        |                            +------------------+   +------------------+
        |                            |     Booking      |   |  RouteWaypoint   |
        |                            |------------------|   |------------------|
        |                            | PK: id           |   | PK: id           |
        |                            | FK: ride_id      |   | FK: ride_id      |
        +--------------------------->| FK: passenger_id |   | sequence_order   |
                                    | seats_requested  |   | latitude, name   |
                                    +------------------+   +------------------+
        """},
        {"type": "p", "bold_prefix": "12.2 Database Schema Tables:", "text": "The tables derived from django models.py are detailed below:"},
        {"type": "table", "headers": ["Table Name", "Columns (Field Name, Type, Constraints)", "Indexes & Relations"],
         "widths": [1.8, 3.2, 2.0],
         "rows": [
             ["core_user", "id (INT AUTO_INCREMENT PK), username (VARCHAR UNIQUE), password (VARCHAR), email (VARCHAR), role (VARCHAR), reputation_points (INT), is_email_verified (BOOL), phone_number (VARCHAR)", "Index on username. Base user settings storage."],
             ["core_vehicle", "id (INT PK), driver_id (INT FK ref core_user.id), make (VARCHAR), model (VARCHAR), license_plate (VARCHAR UNIQUE), capacity (INT), color (VARCHAR), verified (BOOL)", "Index on license_plate. FK driver_id Cascade."],
             ["core_ride", "id (INT PK), driver_id (INT FK ref core_user.id), vehicle_id (INT FK ref core_vehicle.id), start_location (VARCHAR), end_location (VARCHAR), departure_time (DATETIME), price_per_seat (DECIMAL), available_seats (INT), status (VARCHAR), route_map (TEXT)", "Index on status. FK driver_id Cascade, FK vehicle_id Protect."],
             ["core_routewaypoint", "id (INT PK), ride_id (INT FK ref core_ride.id), sequence_order (INT), name (VARCHAR), latitude (DECIMAL), longitude (DECIMAL), estimated_arrival (DATETIME)", "Index on (ride_id, sequence_order) UNIQUE. Sequential route map coordinates storage."],
             ["core_booking", "id (INT PK), ride_id (INT FK ref core_ride.id), passenger_id (INT FK ref core_user.id), pickup_location (VARCHAR), pickup_lat (DECIMAL), dropoff_location (VARCHAR), dropoff_lat (DECIMAL), seats_requested (INT), total_price (DECIMAL), status (VARCHAR)", "FK ride_id Cascade, FK passenger_id Cascade. Booking transaction records."],
             ["core_community", "id (INT PK), name (VARCHAR UNIQUE), description (TEXT), category (VARCHAR), community_type (VARCHAR), invite_link (VARCHAR UNIQUE), created_by_id (INT FK ref core_user.id)", "Index on name. Circles list."],
             ["core_communitymember", "id (INT PK), community_id (INT FK ref core_community.id), user_id (INT FK ref core_user.id), role (VARCHAR), status (VARCHAR)", "Index on (community_id, user_id) UNIQUE. Handles circle authorization."],
             ["core_friendship", "id (INT PK), user_id (INT FK ref core_user.id), friend_id (INT FK ref core_user.id), status (VARCHAR)", "Index on (user_id, friend_id) UNIQUE. Direct friend maps."],
             ["core_message", "id (INT PK), sender_id (INT FK ref core_user.id), recipient_id (INT FK ref core_user.id NULL), community_id (INT FK ref core_community.id NULL), ride_id (INT FK ref core_ride.id NULL), content (TEXT), file_attachment (VARCHAR), is_read (BOOL)", "FK sender_id Cascade. Messaging room buffers."],
             ["core_notification", "id (INT PK), recipient_id (INT FK ref core_user.id), sender_id (INT FK ref core_user.id NULL), notification_type (VARCHAR), content (TEXT), link (VARCHAR), is_read (BOOL)", "FK recipient_id Cascade. Real-time alert triggers."],
             ["core_event", "id (INT PK), title (VARCHAR), location (VARCHAR), date (DATETIME), event_type (VARCHAR), creator_id (INT FK ref core_user.id), community_id (INT FK ref core_community.id NULL)", "FK creator_id Cascade. Event scheduling data."],
             ["core_eventattendee", "id (INT PK), event_id (INT FK ref core_event.id), user_id (INT FK ref core_user.id), status (VARCHAR)", "Index on (event_id, user_id) UNIQUE. Attending registries."],
             ["core_report", "id (INT PK), reporter_id (INT FK ref core_user.id), reported_user_id (INT FK ref core_user.id NULL), reported_ride_id (INT FK ref core_ride.id NULL), reported_community_id (INT FK ref core_community.id NULL), reason (TEXT), status (VARCHAR)", "FK reporter_id Cascade. Safety monitoring logs."],
             ["core_sosalert", "id (INT PK), ride_id (INT FK ref core_ride.id), user_id (INT FK ref core_user.id), latitude (DECIMAL NULL), longitude (DECIMAL NULL), is_active (BOOL)", "FK ride_id Cascade. Emergency tracking alarms."],
             ["core_rating", "id (INT PK), reviewer_id (INT FK ref core_user.id), reviewee_id (INT FK ref core_user.id), ride_id (INT FK ref core_ride.id NULL), rating (INT), comment (TEXT)", "Index on (reviewer_id, reviewee_id, ride_id) UNIQUE. User review ratings."],
             ["core_emaillog", "id (INT PK), recipient (VARCHAR), subject (VARCHAR), content (TEXT), email_type (VARCHAR)", "Read-only logs audit email simulations."],
             ["core_ailog", "id (INT PK), prompt (TEXT), response (TEXT), log_type (VARCHAR)", "Logs prompts/replies details of Gemini integrations."],
             ["core_emailverificationtoken", "id (INT PK), user_id (INT FK ref core_user.id), token (VARCHAR UNIQUE), expires_at (DATETIME), is_used (BOOL)", "FK user_id Cascade. Holds 24h validation strings."],
             ["core_passwordresettoken", "id (INT PK), user_id (INT FK ref core_user.id), token (VARCHAR UNIQUE), expires_at (DATETIME), is_used (BOOL)", "FK user_id Cascade. 1-hour reset tokens database."],
             ["core_communitypost", "id (INT PK), community_id (INT FK ref core_community.id), author_id (INT FK ref core_user.id), content (TEXT), image (VARCHAR NULL), is_announcement (BOOL), likes_count (INT)", "FK community_id Cascade. Feed wall database."],
             ["core_supportticket", "id (INT PK), ticket_number (VARCHAR UNIQUE), name (VARCHAR), email (VARCHAR), subject (VARCHAR), message (TEXT), ai_suggested_reply (TEXT), admin_reply (TEXT), status (VARCHAR), assigned_to_id (INT FK ref core_user.id NULL)", "Index on ticket_number. Help center tickets database."]
         ]},
        {"type": "p", "bold_prefix": "12.3 Normalization Standard: ", "text": "The database design complies with Third Normal Form (3NF). Repeating groups (like waypoint coordinates) are isolated into the RouteWaypoint table. Transitive dependencies are removed by creating dedicated Vehicle and Booking tables, ensuring that a change in vehicle details does not introduce anomalies in active ride records."}
    ]
}

srs_sections["13.0"] = {
    "title": "13. API Documentation",
    "level": 1,
    "blocks": [
        {"type": "p", "text": "This section details the primary backend API routes. All endpoints enforce CSRF checks in session requests and return JSON or HTML redirects as defined:"},
        {"type": "table", "headers": ["Endpoint", "Method", "Auth", "Body/Parameters", "Status Code & Response"],
         "widths": [1.8, 0.8, 1.0, 2.0, 1.9],
         "rows": [
             ["/login/", "POST", "None", "username, password", "302 Redirect to /dashboard/ upon success; 400 Bad Request if fields invalid."],
             ["/register/", "POST", "None", "username, email, password, role, phone_number", "302 Redirect to verification page. Checks username uniqueness."],
             ["/api/check-username/", "GET", "None", "?username=text", "200 OK. JSON: {'available': boolean} (uniqueness checker)."],
             ["/ride/publish/", "POST", "Driver Session", "vehicle, start_location, end_location, departure_time, price_per_seat, available_seats, waypoints (JSON array string)", "302 Redirect to dashboard. RouteWaypoints populated, Ride status set to ACTIVE."],
             ["/ride/search/", "GET", "Passenger Session", "?pickup_lat=decimal&pickup_lng=decimal&dropoff_lat=decimal&dropoff_lng=decimal&seats_needed=int", "200 OK HTML. Renders matched rides sorted by recommendation score (Detour + Trust)."],
             ["/booking/create/<ride_id>/", "POST", "Passenger Session", "seats_requested, pickup_location, dropoff_location", "302 Redirect to payment simulation view. Booking status set to PENDING."],
             ["/booking/<booking_id>/action/", "POST", "Driver Session", "action ('accept' / 'reject')", "200 OK JSON: {'status': 'success'}. Modifies seats capacity, alerts passenger."],
             ["/safety/sos/", "POST", "Session (Active Trip)", "ride_id, latitude, longitude", "200 OK JSON: {'status': 'alert_logged'}. SOSAlert stored, triggers admin alerts."],
             ["/support/chatbot/", "POST", "Session", "question, history (text)", "200 OK JSON: {'answer': text}. Contacts Gemini API support Bot."],
             ["/api/admin/stats/", "GET", "Admin Session", "None", "200 OK JSON. Returns user roles counts, active rides, email volumes, AI log statistics."]
         ]}
    ]
}

srs_sections["14.0"] = {
    "title": "14. AI Module Specifications",
    "level": 1,
    "blocks": [
        {"type": "p", "bold_prefix": "14.1 Purpose: ", "text": "To automate customer care, analyze message safety, categorize incoming service tickets, and compile travel itineraries for carpoolers using Generative AI models."},
        {"type": "p", "bold_prefix": "14.2 Architecture: ", "text": "The Django core application communicates with the Google Gemini API using the google-generativeai SDK. If keys are missing, the ai_services.py module intercepts calls and provides high-fidelity simulated response fallbacks."},
        {"type": "p", "bold_prefix": "14.3 Model: ", "text": "Gemini-1.5-flash is configured as the default model. Temperature settings vary: 0.2 for strict classifications (e.g. support emails, toxicity analyses), and 0.7 for creative text generations (travel itineraries)."},
        {"type": "p", "bold_prefix": "14.4 Prompt Flow: ", "text": "1. User inputs text. 2. Backend formats request using prompt templates. 3. System checks toxicity. 4. Executes Gemini API call. 5. Logs request and response parameters in AILog table. 6. Validates format. 7. Renders to user."},
        {"type": "image", "name": "ai_chatbot.png", "text": "Figure 14.1: Helpdesk Automated Support Chatbot Session"},
        {"type": "image", "name": "ai_trip_planner.png", "text": "Figure 14.2: AI-Generated Travel Route Itinerary Assistant"},
        {"type": "p", "bold_prefix": "14.5 Fallback, Safety, & Token Limits: ", "text": "A custom toxicity check flags harassment. If text contains scam words, is_toxic=True is saved, and content is blocked. Caching is simulated, and Google's free-tier rate limits (15 Requests Per Minute) are handled with a fallback mechanism that switches to local regex templates if limits are breached. Error handling logs API issues in standard Django files without throwing 500 errors to the client."},
        {"type": "p", "bold_prefix": "14.6 Prompt Templates in Code:", "text": "The exact prompts defined in core/ai_services.py are:"},
        {"type": "code", "text": """
# Template 1: AI Support Chatbot Context
Context: You are a helpful 24/7 Support Assistant for "Ride-Share", a real-time carpooling and community travel platform.
Chat History: {chat_history}
User Question: {user_question}
Answer the user's question clearly, concisely, and professionally. Explain the rules, platform features (like communities, SOS alarms, payment sharing), and support channels if they need user verification.

# Template 2: Support Ticket & Email Auto-Reply Classifier
Read the following support email:
Sender: {sender_email}
Subject: {subject}
Body: {content}
Perform intent detection, classify the request, and draft a professional reply.
Return ONLY a JSON object with the following keys:
1. "category": Choose from (Account Issue, Login Problem, Verification Issue, Travel Question, Report Abuse, Community Support, Payment Query, Feature Request).
2. "needs_escalation": boolean (True if it involves safety reports, billing fraud, or technical bugs).
3. "reply_content": Markdown text of your drafted reply to the user.
        """}
    ]
}

srs_sections["15.0"] = {
    "title": "15. Email System",
    "level": 1,
    "blocks": [
        {"type": "p", "text": "The platform requires robust automated emails to manage password security and account verification:"},
        {"type": "bullet", "bold_prefix": "SMTP Configuration: ", "text": "Configured in settings.py using EMAIL_BACKEND. In production, connects to secure hosts (SendGrid, Mailgun) using SSL/TLS via port 465. In development, writes outputs directly to DB logs (EmailLog model) to prevent unsolicited spam transmissions."},
        {"type": "bullet", "bold_prefix": "Verification Flow: ", "text": "Upon registration, a 64-character token is saved (expires in 24 hours). The simulated verification link routes to /verify-email/?token=value. Entering the token changes User.is_email_verified to True."},
        {"type": "bullet", "bold_prefix": "Password Reset: ", "text": "Password recovery triggers PasswordResetToken creation (expires in 1 hour). The link routes to /reset-password/?token=value, verifying expiration before updating password hashes."},
        {"type": "bullet", "bold_prefix": "AI Auto-Reply Simulation: ", "text": "Admins can process incoming support emails. The system calls Gemini to classify details, draft responses, and queue them in EmailLog logs automatically."},
        {"type": "bullet", "bold_prefix": "Spam Protection: ", "text": "Rate limits email verification triggers to a maximum of 3 resends per hour per email address to defend against SMTP pool depletion."}
    ]
}

srs_sections["16.0"] = {
    "title": "16. Security Infrastructure",
    "level": 1,
    "blocks": [
        {"type": "p", "text": "Security is built into the Django backend using secure defaults and custom rate-limit modules:"},
        {"type": "bullet", "bold_prefix": "Authentication & Session Management: ", "text": "Standard sessions are maintained via signed cookies. Session identifiers are rotated upon login to defend against session fixation vulnerabilities. AJAX queries utilize CSRF token verification flags."},
        {"type": "bullet", "bold_prefix": "Authorization controls: ", "text": "Access checks are enforced at database and view layer. Routes verify ownership before executing edits (e.g. a driver cannot delete another driver's ride)."},
        {"type": "bullet", "bold_prefix": "Password Hashing: ", "text": "Uses Django's default PBKDF2 algorithm with SHA-256 and 800,000 iterations to protect passwords in the database."},
        {"type": "bullet", "bold_prefix": "Injection Protections: ", "text": "Django's ORM parameters query building automatically protects database calls from SQL Injection. Front-end HTML templates use engine escaping constraints to protect against Cross-Site Scripting (XSS)."},
        {"type": "bullet", "bold_prefix": "Rate Limiting: ", "text": "django-ratelimit handles rate limits on search queries. Login views utilize django-axes to lock profiles after 5 consecutive authentication failures from a single IP within 5 minutes."},
        {"type": "bullet", "bold_prefix": "Secrets Management: ", "text": "Secrets (Gemini API keys, database credentials, Django SECRET_KEY) are loaded from .env file configurations using python-decouple. Hardcoded credentials in source control are strictly prohibited."}
    ]
}

srs_sections["17.0"] = {
    "title": "17. UI/UX Design System",
    "level": 1,
    "blocks": [
        {"type": "p", "text": "The front-end user experience emphasizes visual clarity, accessibility compliance, and mobile responsive controls:"},
        {"type": "bullet", "bold_prefix": "Design Principles: ", "text": "Clean grid alignment, high typography hierarchies, visual icons cues (Bootstrap Icons), and clear interactive states."},
        {"type": "image", "name": "home_page.png", "text": "Figure 17.1: Desktop Landing/Home Page Canvas"},
        {"type": "bullet", "bold_prefix": "Color System: ", "text": "Defined in static/css/style.css using CSS Custom Properties for seamless Light/Dark mode toggling. Primary brand color: Deep Navy Blue (#003366). Secondary accents: Soft Gray (#E0E0E0), Success state: Forest Green (#28A745), Warning alarm: Crimson Red (#DC3545)."},
        {"type": "bullet", "bold_prefix": "Typography: ", "text": "Google Font 'Outfit' or 'Inter' is loaded as the primary font family. System fallbacks include sans-serif. Heading weights are bold (700); body paragraphs are regular (400)."},
        {"type": "bullet", "bold_prefix": "Responsive Design: ", "text": "Adapts to desktop, tablet, and mobile displays using Bootstrap 5 fluid grid containers. Key CSS breakpoints: Mobile (<576px), Tablet (<768px), Desktop (>=992px)."},
        {"type": "image", "name": "mobile_dashboard.png", "text": "Figure 17.2: Mobile Responsive Passenger Dashboard Interface Layout (375x812 View)"},
        {"type": "bullet", "bold_prefix": "Micro-Animations & Visuals: ", "text": "Includes smooth hover transitions (0.2s ease-in-out) on dashboard cards and form submit buttons, alongside a pulsing animation on the active red SOS button."},
        {"type": "image", "name": "error_page.png", "text": "Figure 17.3: Custom 404 URL Not Found Error State"}
    ]
}

srs_sections["18.0"] = {
    "title": "18. External Integrations",
    "level": 1,
    "blocks": [
        {"type": "p", "text": "The platform relies on several key external integrations to function:"},
        {"type": "table", "headers": ["External Service", "Purpose", "Authentication", "Limits & Quotas", "Fallback Strategy"],
         "widths": [1.5, 1.8, 1.2, 1.2, 1.8],
         "rows": [
             ["Leaflet.js / OSM", "Loads interactive maps and plots sequential waypoint route markers.", "None (Free OpenStreetMap Tile layers).", "Unlimited tile loading requests.", "Displays warning banner; map reverts to standard inputs form fields."],
             ["Nominatim API", "Geocodes address strings into GPS coordinates (latitude, longitude).", "None (Relies on User Agent identifiers).", "Max 1 request per second.", "Displays address lookup errors; requests users coordinate inputs manually."],
             ["OSRM API", "Calculates route distances and sequential waypoints routing paths.", "None (Public routing endpoint).", "Rate limits on shared servers.", "Plots straight geodetic routes based on mathematical Haversine sequence."],
             ["Google Gemini API", "Powers helpdesk support chatbot, support email classifiers, and travel itinerary builders.", "Bearer API Key (GEMINI_API_KEY env).", "15 Requests Per Minute (RPM) on free tier.", "Intercepts requests in ai_services.py, fallback to local regex drafts and mock templates."],
             ["SMTP Relay Gate", "Dispatches password recovery and account verification notifications.", "Credentials Login (Host, Port, User, Password).", "Varies by provider (SendGrid free: 100/day).", "Logs email body in DB (EmailLog) for manual operator recovery."]
         ]}
    ]
}

srs_sections["19.0"] = {
    "title": "19. Deployment Guide",
    "level": 1,
    "blocks": [
        {"type": "p", "text": "Deployment environments and configuration requirements are defined as follows:"},
        {"type": "bullet", "bold_prefix": "Environments: ", "text": "Development: Runs Django debug local server (127.0.0.1:8000), using SQLite database db.sqlite3. Production: Runs in Docker containers on Render, connected to PostgreSQL instances with debug turned off (DEBUG=False)."},
        {"type": "bullet", "bold_prefix": "CI/CD Pipeline: ", "text": "GitHub Actions automatically execute on all push requests to main branch. The pipeline boots python venv, installs requirements.txt, runs migrations checks, executes unit test suites, and alerts of build failures."},
        {"type": "bullet", "bold_prefix": "Git Workflow: ", "text": "Follows standard GitFlow practices. Features are developed in separate feature/ branches and merged into master/main via verified pull requests."},
        {"type": "bullet", "bold_prefix": "Server Config (Render/Vercel): ", "text": "Renders configuration via render.yaml (packages ASGI daphne/gunicorn processes). Vercel handles static assets deployments via vercel.json configurations. SSL certs are auto-renewed via Let's Encrypt."},
        {"type": "bullet", "bold_prefix": "Environment Variables: ", "text": "The .env file must store SECRET_KEY, GEMINI_API_KEY, DB_URL, EMAIL_HOST_USER, and EMAIL_HOST_PASSWORD. Static files are collected using collectstatic and served via Whitenoise."}
    ]
}

srs_sections["20.0"] = {
    "title": "20. Testing & Verification",
    "level": 1,
    "blocks": [
        {"type": "p", "text": "The quality assurance framework verifies that all functional and security requirements are met:"},
        {"type": "bullet", "bold_prefix": "Unit Testing: ", "text": "Written in backend/core/tests.py using Django's TestCase. The suite covers: User creation, geodetic waypoint matching validations, booking status updates, vehicle capacity validations, and API logins."},
        {"type": "bullet", "bold_prefix": "Integration Testing: ", "text": "Verifies interactions between route matching and booking creations (e.g. Booking seat allocation updates active ride listings)."},
        {"type": "bullet", "bold_prefix": "System & UAT Testing: ", "text": "Manual verification checklists run inside staging environments, simulating passengers booking rides, drivers reviewing lists, and clicking the red SOS panic alarm system."},
        {"type": "bullet", "bold_prefix": "Performance testing: ", "text": "Simulates up to 10,000 active routes in database seeding, verifying that page load times and search latencies remain within NFR targets."},
        {"type": "bullet", "bold_prefix": "Test Cases Checklist:", "text": ""},
        {"type": "table", "headers": ["Test ID", "Description", "Inputs", "Expected Outcome"],
         "widths": [1.0, 2.5, 1.5, 2.5],
         "rows": [
             ["TC-001", "Verify new passenger account registration.", "Username='testuser', Role='PASSENGER'", "User record created, inactive state. Verification token saved in DB."],
             ["TC-002", "Check password recovery token validation.", "Valid email address.", "PasswordResetToken generated, expires in 1 hour. Email log recorded."],
             ["TC-003", "Intelligent Geodetic Sequential Matching Check.", "Pickup coordinates matching active ride sequential waypoints.", "Ride returned. Walk distance <= maximum walk distance. Ranks by recommendation score."],
             ["TC-004", "Prevent reverse sequence route bookings.", "Pickup point located after dropoff point along driver route.", "Route matching filters ride out of recommendation list (Pickup Index > Dropoff Index)."],
             ["TC-005", "Trigger SOS Panic Alert.", "Active trip user triggers SOS.", "SOSAlert stored in DB as active. Siren audio element plays on client UI."],
             ["TC-006", "Check toxicity content filter.", "Hate speech message text content.", "AI filter flags toxicity. returns is_toxic=True, blocks message from feeds."]
         ]}
    ]
}

srs_sections["21.0"] = {
    "title": "21. Risks Assessment",
    "level": 1,
    "blocks": [
        {"type": "p", "text": "Key operational, legal, and technical risks are defined below:"},
        {"type": "bullet", "bold_prefix": "Technical: ", "text": "SQLite database concurrency bottlenecks. Mitigation: Migrate to PostgreSQL in production environments."},
        {"type": "bullet", "bold_prefix": "Business: ", "text": "Lack of driver liquidity or low passenger density. Mitigation: Target specific corporate campuses and colleges to establish high-density routes first."},
        {"type": "bullet", "bold_prefix": "Legal: ", "text": "Compliance liabilities with local commercial transport regulators. Mitigation: Set strict limits on seat pricing to ensure drivers only offset fuel costs without making profit."},
        {"type": "bullet", "bold_prefix": "Security: ", "text": "Threat of harassment or fake user accounts. Mitigation: Require government identity uploads and corporate/college email verification badges."},
        {"type": "bullet", "bold_prefix": "AI Risks: ", "text": "Gemini API rate limiting or hallucinations in FAQ support answers. Mitigation: Set temperature to 0.2, implement fallback mock answers, and allow ticket escalation for human review."},
        {"type": "bullet", "bold_prefix": "Deployment: ", "text": "Server hosting outages (Render/Vercel). Mitigation: Run redundant mirrors and host PWA service workers caching policies for offline-first dashboard accessibility."}
    ]
}

srs_sections["22.0"] = {
    "title": "22. Future Enhancements",
    "level": 1,
    "blocks": [
        {"type": "p", "text": "The roadmap for releases after MVP 1.0 includes the following improvements:"},
        {"type": "bullet", "bold_prefix": "1. Live GPS tracking: ", "text": "Integrate real-time geolocation tracking using HTML5 Geolocation API, WebSockets (Django Channels), and background device services."},
        {"type": "bullet", "bold_prefix": "2. Production Payment Gateway: ", "text": "Transition from simulated checkout to a live payment gateway integration (Stripe, Razorpay, or PayPal API)."},
        {"type": "bullet", "bold_prefix": "3. Native Mobile Compilation: ", "text": "Package the PWA codebase into native Android and iOS wrappers using Capacitor or Apache Cordova frameworks to publish to official App Stores."},
        {"type": "bullet", "bold_prefix": "4. Multi-Modal Travel Matching: ", "text": "Expand recommendation engine to combine public transit data (bus, subway routes) with carpooling legs for long-distance commutes."}
    ]
}

srs_sections["23.0"] = {
    "title": "23. Maintenance Strategy",
    "level": 1,
    "blocks": [
        {"type": "p", "text": "System maintenance procedures are defined to guarantee long-term operational health:"},
        {"type": "bullet", "bold_prefix": "Bug Fixes & Patching: ", "text": "Urgent security patches (e.g. Django security updates) are deployed within 24 hours of release. Minor bugs are compiled into monthly release packages."},
        {"type": "bullet", "bold_prefix": "Updates: ", "text": "Dependencies in requirements.txt are reviewed quarterly. Minor and patch updates are automated using package managers (e.g. Dependabot)."},
        {"type": "bullet", "bold_prefix": "Versioning: ", "text": "Strict adherence to Semantic Versioning (SemVer) guidelines: MAJOR.MINOR.PATCH. Pre-releases are tagged as -beta or -rc."},
        {"type": "bullet", "bold_prefix": "Incident Response: ", "text": "If a server crash or database outage occurs, automated pager alerts trigger. The on-call DevOps engineer initiates rollback scripts, verifies logs, and reports root-cause analyses (RCA) in support archives."}
    ]
}

srs_sections["24.0"] = {
    "title": "24. Project Timeline",
    "level": 1,
    "blocks": [
        {"type": "p", "text": "The timeline spans a 24-week lifecycle from initial requirements gathering to final production release:"},
        {"type": "code", "text": """
Phases & Weeks Tracker:
  W1-W4   : Planning & SRS Documentation [Completed]
  W5-W8   : Database Setup & Geodetic Logic Coding [Completed]
  W9-W14  : Frontend Integration, Map Editor & Chat Rooms [Completed]
  W15-W18 : AI Helper bots & Email Systems Integration [Completed]
  W19-W20 : QA Testing & Security Audits [In Progress]
  W21-W22 : Staging deploy, UAT Vetting [Upcoming]
  W23-W24 : Production Release & Post-deployment reviews [Upcoming]
        """},
        {"type": "table", "headers": ["Milestone", "Target Week", "Deliverables", "Status"],
         "widths": [1.8, 1.0, 3.0, 1.2],
         "rows": [
             ["M1: Specification Approval", "Week 4", "Software Requirements Specification (SRS) Document.", "Completed"],
             ["M2: Backend Baseline", "Week 8", "Django app models structure, database seeding, unit tests.", "Completed"],
             ["M3: UI Integration", "Week 14", "Leaflet waypoint sequencing interface, responsive dashboard templates.", "Completed"],
             ["M4: System Completion", "Week 18", "Gemini Support bots, simulated SMTP logs, chat uploads, offline cache.", "Completed"],
             ["M5: Production Release", "Week 24", "Render hosting deploy, production DB migrations, security checks.", "Upcoming"]
         ]}
    ]
}

srs_sections["25.0"] = {
    "title": "25. Budget & Cost Breakdown",
    "level": 1,
    "blocks": [
        {"type": "p", "text": "Project financial projections and operational billing costs are defined as follow:"},
        {"type": "bullet", "bold_prefix": "Development Cost: ", "text": "=== FILLED BY USER ==="},
        {"type": "bullet", "bold_prefix": "Hosting Servers (Render/AWS): ", "text": "=== FILLED BY USER ==="},
        {"type": "bullet", "bold_prefix": "Database Cloud Instance: ", "text": "=== FILLED BY USER ==="},
        {"type": "bullet", "bold_prefix": "Domain Name Registration: ", "text": "=== FILLED BY USER ==="},
        {"type": "bullet", "bold_prefix": "Email Provider API (SendGrid): ", "text": "=== FILLED BY USER ==="},
        {"type": "bullet", "bold_prefix": "Google Gemini API Usage: ", "text": "=== FILLED BY USER ==="},
        {"type": "bullet", "bold_prefix": "Licensing Fees: ", "text": "=== FILLED BY USER === (System utilizes free MIT/BSD open source libraries, resulting in ₹0 licensing fees)."},
        {"type": "bullet", "bold_prefix": "Monthly Maintenance Billing: ", "text": "=== FILLED BY USER ==="}
    ]
}

srs_sections["26.0"] = {
    "title": "26. Legal & Compliance",
    "level": 1,
    "blocks": [
        {"type": "p", "text": "To satisfy university standards, startup governance, or industrial deployments, the system must address regulatory requirements:"},
        {"type": "bullet", "bold_prefix": "Privacy Policy: ", "text": "=== FILLED BY USER ===. The policy must cover user data retention, coordinates logging histories, and ID vetting storage policies."},
        {"type": "bullet", "bold_prefix": "Terms of Service: ", "text": "=== FILLED BY USER ===. Must declare that Ride-Share is a peer-to-peer cost-sharing platform and is not a commercial taxi aggregation network."},
        {"type": "bullet", "bold_prefix": "Cookie Policy: ", "text": "=== FILLED BY USER ===. Session and CSRF cookies are strictly essential for core authentication operations; no marketing cookies are utilized."},
        {"type": "bullet", "bold_prefix": "GDPR Compliance: ", "text": "Allows users to execute 'Right to be Forgotten' commands in profile settings, executing cascading deletion of usernames, credentials, and location records from DB storage."},
        {"type": "bullet", "bold_prefix": "Applicable Jurisdiction Laws: ", "text": "=== FILLED BY USER ==="},
        {"type": "bullet", "bold_prefix": "Copyright Ownership: ", "text": "Copyright 2026 Ride-Share Development Team. All rights reserved. Intellectual property rights belong to === FILLED BY USER ===."}
    ]
}

srs_sections["27.0"] = {
    "title": "27. Appendices",
    "level": 1,
    "blocks": [
        {"type": "p", "bold_prefix": "Glossary of terms: ", "text": "1. Haversine: An equation giving great-circle distances between two pairs of coordinates on a sphere. 2. Leaflet.js: A lightweight, mobile-friendly Javascript mapping library. 3. SQLite: A lightweight serverless SQL database engine. 4. PWA: Progressive Web App capability installable on client OS layers."},
        {"type": "p", "bold_prefix": "System Configuration Template (settings.py environment variables):", "text": ""},
        {"type": "code", "text": """
# .env File Configuration Blueprint
SECRET_KEY=django-insecure-mock-key-development-only
DEBUG=True
ALLOWED_HOSTS=localhost,127.0.0.1,*.render.com
DATABASE_URL=sqlite:///db.sqlite3
GEMINI_API_KEY=AIzaSyYourGeminiApiKeyHere
EMAIL_BACKEND=django.core.mail.backends.console.EmailBackend
EMAIL_HOST=smtp.sendgrid.net
EMAIL_PORT=587
EMAIL_USE_TLS=True
EMAIL_HOST_USER=apikey
EMAIL_HOST_PASSWORD=SendGridApiPasswordGoesHere
        """},
        {"type": "p", "bold_prefix": "Deployment Checklist: ", "text": "1. Set DEBUG=False. 2. Collect static files (python manage.py collectstatic). 3. Run migrations checks. 4. Setup SSL/TLS Let's Encrypt certificates. 5. Populate environment keys in Render. 6. Verify Service Worker sw.js matches route patterns."},
        {"type": "p", "bold_prefix": "Release Checklist: ", "text": "1. Run full unit testing suites. 2. Verify PWA install icon sizes (192px/512px). 3. Verify clean seeding (python manage.py seed_demo). 4. Perform accessibility audit scans."},
        {"type": "p", "bold_prefix": "Support Contact Directory: ", "text": "Technical lead: support@yourdomain.com. System Admin emergency helpline: === FILLED BY USER ===."}
    ]
}

def generate_markdown(md_sections, filepath):
    """Generates the Markdown version of the SRS containing links to screenshots."""
    with open(filepath, "w", encoding="utf-8") as f:
        # Title
        f.write("# Software Requirements Specification (SRS)\n\n")
        f.write("**Project Name:** AI-Based Smart Route Ride-Sharing System  \n")
        f.write("**Product Version:** 1.0.0 (Production-Inspired MVP)  \n")
        f.write("**Date:** July 7, 2026  \n")
        f.write("**Document Status:** Approved for Implementation  \n")
        f.write("**Prepared By:** === FILLED BY USER ===  \n")
        f.write("**Organization:** === FILLED BY USER ===  \n")
        f.write("**Client:** === FILLED BY USER ===  \n")
        f.write("**Confidentiality Statement:** === FILLED BY USER ===  \n")
        f.write("**Document Classification:** === FILLED BY USER ===  \n\n")
        f.write("---\n\n")
        
        # Document Control Table of Contents etc.
        f.write("## Document Control\n\n")
        f.write("### Revision History\n")
        f.write("| Ver. | Date | Author | Reviewer(s) | Description of Change |\n")
        f.write("|---|---|---|---|---|\n")
        f.write("| 0.1.0 | 2026-07-01 | === FILLED BY USER === | === FILLED BY USER === | Initial draft outline & requirements gathering. |\n")
        f.write("| 1.0.0 | 2026-07-07 | === FILLED BY USER === | === FILLED BY USER === | Baseline release mapping fully to source models and features. |\n\n")
        
        f.write("### Approvals\n")
        f.write("| Name | Role | Signature | Date |\n")
        f.write("|---|---|---|---|\n")
        f.write("| === FILLED BY USER === | === FILLED BY USER === | === FILLED BY USER === | === FILLED BY USER === |\n\n")
        
        f.write("### Distribution List\n")
        f.write("| Name / Group | Organization / Role | Date Shared |\n")
        f.write("|---|---|---|\n")
        f.write("| === FILLED BY USER === | === FILLED BY USER === | === FILLED BY USER === |\n\n")
        
        f.write("---\n\n")
        f.write("## Table of Contents\n\n")
        for key, sec in md_sections.items():
            lvl = sec.get('level', 1)
            title = sec.get('title', key)
            anchor = title.lower().replace('.', '').replace(' ', '-').replace('&', '').replace('(', '').replace(')', '').replace(':', '')
            f.write("  " * (lvl - 1) + f"- [{title}](#{anchor})\n")
        f.write("\n---\n\n")
        
        # Write Sections
        for key, sec in md_sections.items():
            lvl = sec.get('level', 1)
            title = sec.get('title', key)
            h_prefix = "#" * (lvl + 1)
            f.write(f"{h_prefix} {title}\n\n")
            
            blocks = sec.get('blocks', [])
            for block in blocks:
                b_type = block.get('type', 'p')
                text = block.get('text', '')
                bold_prefix = block.get('bold_prefix', None)
                
                if b_type == 'p':
                    if bold_prefix:
                        f.write(f"**{bold_prefix}** {text}\n\n")
                    else:
                        f.write(f"{text}\n\n")
                elif b_type == 'bullet':
                    if bold_prefix:
                        f.write(f"* **{bold_prefix}** {text}\n")
                    else:
                        f.write(f"* {text}\n")
                elif b_type == 'code':
                    f.write(f"```\n{text}\n```\n\n")
                elif b_type == 'image':
                    img_name = block.get('name')
                    f.write(f"![{text}](file:///d:/Ride-Share/screenshots/{img_name})\n")
                    f.write(f"*{text}*\n\n")
                elif b_type == 'table':
                    # Table formatting
                    headers = block.get('headers', [])
                    rows = block.get('rows', [])
                    f.write("| " + " | ".join(headers) + " |\n")
                    f.write("| " + " | ".join(["---"] * len(headers)) + " |\n")
                    for row in rows:
                        f.write("| " + " | ".join([str(val) for val in row]) + " |\n")
                    f.write("\n")
            if any(b.get('type') == 'bullet' for b in blocks):
                f.write("\n") # Add spacer after bullet list
                
    print(f"Markdown document saved to {filepath}")

if __name__ == "__main__":
    doc_path = r"d:\Ride-Share\docs\05_SRS\SOFTWARE_REQUIREMENTS_SPECIFICATION.docx"
    md_path = r"d:\Ride-Share\docs\05_SRS\SOFTWARE_REQUIREMENTS_SPECIFICATION.md"
    
    # Generate both formats
    generate_markdown(srs_sections, md_path)
    build_docx(srs_sections, doc_path)
